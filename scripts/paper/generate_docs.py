import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        r = table.add_row()
        for i, val in enumerate(row):
            r.cells[i].text = val
    doc.add_paragraph('')
    return table

# ================================================================
# TITLE
# ================================================================
p = doc.add_heading('VANI - Voice Analysis & Neural Intelligence', 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph('Military-Grade Offline Radio Intercept Analysis System (SIGINT)')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(13)

p3 = doc.add_paragraph('System Documentation: Models & Application Tabs')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

p4 = doc.add_paragraph('Date: 2026-03-14 | Path: C:/offline_ai_system_v2/')
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ================================================================
# SECTION 1: OVERVIEW
# ================================================================
add_heading(doc, '1. Project Overview', 1)
doc.add_paragraph(
    'VANI (Voice Analysis & Neural Intelligence) is a fully offline, CPU-only military-grade '
    'radio intercept analysis system designed to run on an 8 GB RAM machine. It processes raw '
    'audio intercepts through a multi-stage AI pipeline and produces structured intelligence '
    'reports (ISUM) in a Streamlit web interface.'
)
doc.add_paragraph('Entry Point: streamlit run app.py')
doc.add_paragraph('Working Directory: C:/offline_ai_system_v2/')

# ================================================================
# SECTION 2: AI MODELS
# ================================================================
add_heading(doc, '2. Installed AI Models', 1)
doc.add_paragraph(
    'All models are stored locally under the models/ directory. No internet connection is required at runtime.'
)

add_table(doc,
    ['Model Directory', 'Purpose', 'Notes'],
    [
        ('whisper-large-v3-turbo-ct2', 'PRIMARY ASR (Speech-to-Text)',
         'CTranslate2 optimised, faster-whisper library. Active transcription model.'),
        ('whisper-large-v3-turbo', 'Alternate Whisper (non-CT2)',
         'Fallback; standard HuggingFace format.'),
        ('whisper_medium', 'Original ASR model',
         'No longer primary; kept for reference.'),
        ('nllb-200-distilled-600M', 'PRIMARY Translation Model',
         'Supports 200 languages including all major Indic languages (Hindi, Punjabi, Urdu, Pashto, etc.).'),
        ('indictrans2-indic-en-1B', 'Secondary Translation (Dogri only)',
         'Used exclusively for Dogri (doi) - the only Indic language not covered by NLLB-200.'),
        ('langid/lid.176.bin', 'Text Language Identification',
         'FastText model - identifies language from transcribed text (176 languages).'),
        ('mms-lid-256', 'Audio Language Identification',
         'Facebook MMS model - identifies language directly from raw audio (256 languages).'),
        ('qwen2.5-0.5b-instruct', 'LLM for ISUM (smaller)',
         'Qwen 0.5B - lighter option for intelligence summary generation.'),
        ('qwen2.5-1.5b-instruct', 'LLM for ISUM (primary)',
         'Qwen 1.5B - active LLM per config.yaml. Auto-activated if present.'),
    ]
)

# 2.1 ASR
add_heading(doc, '2.1  ASR - Automatic Speech Recognition', 2)
doc.add_paragraph(
    'Model: whisper-large-v3-turbo-ct2 (CTranslate2 format, faster-whisper library)\n'
    '\n'
    'Function: Converts audio speech to text (transcript).\n'
    '\n'
    'Configuration:\n'
    '  - beam_size = 4, temperature = 0.0\n'
    '  - condition_on_previous_text = false\n'
    '  - Segments with no_speech_prob > 0.70 are discarded (silence/noise filtering)\n'
    '  - Language detected per-chunk; result cached after first chunk for speed\n'
    '  - Initial prompt includes Punjabi Gurmukhi phrases to improve recognition accuracy for border-region intercepts'
)

# 2.2 LangID
add_heading(doc, '2.2  Language Identification (LangID) - 3-Way Voting System', 2)
doc.add_paragraph(
    'VANI uses three independent models to determine the spoken language, then combines their '
    'results through a confidence-weighted voting system.\n'
)

add_table(doc,
    ['Model', 'Input', 'Coverage'],
    [
        ('Whisper Language Prob', 'Audio waveform (via ASR)', 'Languages supported by Whisper'),
        ('FastText (lid.176.bin)', 'Transcribed text', '176 languages'),
        ('MMS-LID-256', 'Raw audio waveform', '256 languages'),
    ]
)

doc.add_paragraph(
    'Voting Logic:\n'
    '  - Unanimous agreement -> highest confidence result selected\n'
    '  - Majority (2 of 3) -> average confidence of agreeing models\n'
    '  - All disagree -> single best-confidence result selected\n'
    '  - If final confidence < 0.60 -> flagged as uncertain in the ISUM report\n'
    '\n'
    'Special Punjabi Fix (Critical):\n'
    'Whisper frequently misidentifies Punjabi as Hindi. If Gurmukhi script is detected '
    'in the transcript, OR if FastText/MMS both return "pa" while Whisper returns "hi", '
    'the system forces the language to Punjabi (pa) and routes through NLLB translation.'
)

# 2.3 Translation
add_heading(doc, '2.3  Translation Models', 2)
doc.add_paragraph(
    'NLLB-200 (nllb-200-distilled-600M) - PRIMARY\n'
    '\n'
    'Translates all non-English, non-Dogri languages to English.\n'
    'Confirmed working for: Hindi (hi), Punjabi (pa), Urdu (ur), Pashto (ps), '
    'Nepali (ne), Bengali (bn), Maithili (mai), Kashmiri (ks), Sindhi (sd), '
    'Sinhala (si), Chinese (zh), Burmese (my), Tibetan (bo), Persian (fa), '
    'Arabic (ar), Tajik (tg), Uzbek (uz), Kazakh (kk).\n'
    '\n'
    'IndicTrans2 (indictrans2-indic-en-1B) - SECONDARY (Dogri only)\n'
    '\n'
    'Used exclusively for Dogri (doi) - the gap language not supported by NLLB-200.\n'
    'Requires use_cache=False and attn_implementation="eager" due to transformer 5.3.0 compatibility.'
)

# 2.4 LLM
add_heading(doc, '2.4  LLM - Intelligence Summary Generator', 2)
doc.add_paragraph(
    'Model: qwen2.5-1.5b-instruct (primary) / qwen2.5-0.5b-instruct (lighter fallback)\n'
    '\n'
    'Function: Generates structured ISUM (Intelligence Summary) reports from translated text.\n'
    '\n'
    'Behaviour:\n'
    '  - Auto-activated if the model directory exists on disk\n'
    '  - Falls back to rule-based extraction if LLM is unavailable (always works)\n'
    '  - Rule-based mode extracts: callsigns (Alpha/Bravo patterns), directions, grid references, time references\n'
    '  - Assigns threat levels: CRITICAL / HIGH / MEDIUM / LOW / CLEAR'
)

# ================================================================
# SECTION 3: PIPELINE
# ================================================================
add_heading(doc, '3. Processing Pipeline (src/pipeline.py)', 1)
doc.add_paragraph(
    'The pipeline processes an audio file through 10 sequential stages. '
    'Output is saved to: output/{audio_stem}_result.json'
)

add_table(doc,
    ['Stage', 'Name', 'Function'],
    [
        ('Stage 1', 'VAD (Voice Activity Detection)',
         'Detects speech segments. Filters out silence and non-speech regions.'),
        ('Stage 2', 'Preprocessing',
         'Normalises audio format, sample rate, and channel count for downstream models.'),
        ('Stage 3', 'Chunking',
         'Splits long audio into manageable chunks aligned to VAD segment boundaries.'),
        ('Stage 4', 'ASR (Transcription)',
         'Whisper converts speech chunks to text. Returns segments with timestamps and confidence scores.'),
        ('Stage 5', 'Language Identification',
         '3-way vote (Whisper + FastText + MMS-LID) determines the spoken language.'),
        ('Stage 6', 'Translation',
         'Routes to NLLB-200 or IndicTrans2 based on detected language. English audio skips this stage.'),
        ('Stage 7', 'Keyword Detection',
         'Scans transcript/translation for threat-related keywords across configured categories.'),
        ('Stage 8', 'ISUM Generation',
         'Qwen LLM (or rule-based fallback) produces structured 5W intelligence summary.'),
    ]
)

# ================================================================
# SECTION 4: TABS
# ================================================================
add_heading(doc, '4. Application Tabs (app.py)', 1)
doc.add_paragraph(
    'The Streamlit interface contains 8 tabs, each with a keyboard shortcut shown in brackets.'
)

add_table(doc,
    ['Tab', 'Shortcut', 'One-Line Purpose'],
    [
        ('PROCESS', '[P]', 'Upload and process audio intercepts through the full pipeline.'),
        ('ISUM REPORT', '[I]', 'View the structured intelligence summary for the last processed file.'),
        ('SEARCH', '[S]', 'Search historical intercepts by keyword, language, threat level, date.'),
        ('DASHBOARD', '[D]', 'Visual analytics: threat distribution, language breakdown, trends.'),
        ('HISTORY', '[H]', 'Browse all past intercepts stored in the local database.'),
        ('EXPORT', '[E]', 'Download the report as PDF, DOCX, or raw JSON.'),
        ('METRICS', '[M]', 'Evaluate ASR and translation quality (Tier 1 auto + Tier 2 manual).'),
        ('ANNOTATE', '[A]', 'Analyst corrections for transcript/translation/5W - builds training data.'),
    ]
)

# Tab details
add_heading(doc, '4.1  [P] PROCESS Tab', 2)
doc.add_paragraph(
    'The main entry point of the system.\n'
    '\n'
    'What you can do:\n'
    '  - Upload an audio file (WAV, MP3, OGG, FLAC, etc.)\n'
    '  - Click Process to run the full 10-stage pipeline\n'
    '  - View results: detected language, transcript, translation, keyword alerts, threat level\n'
    '\n'
    'What happens behind the scenes:\n'
    '  - Result stored in session state (st.session_state["last_result"])\n'
    '  - Result saved to output/{stem}_result.json\n'
    '  - Result saved to the SQLite database\n'
    '  - Processing time and memory usage displayed after completion'
)

add_heading(doc, '4.2  [I] ISUM REPORT Tab', 2)
doc.add_paragraph(
    'Displays the structured Intelligence Summary (ISUM) report.\n'
    '\n'
    'Contents:\n'
    '  - Full 5W breakdown: Who, What, Where, When, Assessment\n'
    '  - Threat level badge (CRITICAL / HIGH / MEDIUM / LOW / CLEAR)\n'
    '  - Confidence flags: LOW_LANG_CONFIDENCE, TRANSLATION_UNRELIABLE, ASR_LOW_CONFIDENCE, TRANSLATION_FAILED\n'
    '  - Keyword alerts and top threat categories\n'
    '\n'
    'Fallback behaviour:\n'
    'If session state is empty (e.g. after app restart), scans output/*_result.json '
    'to load the most recently processed file automatically.'
)

add_heading(doc, '4.3  [S] SEARCH Tab', 2)
doc.add_paragraph(
    'Full-text and filtered search across all stored intercepts.\n'
    '\n'
    'Search filters available:\n'
    '  - Keyword/phrase (full-text search in transcript and translation)\n'
    '  - Language code (e.g. hi, pa, ur)\n'
    '  - Threat level (CRITICAL, HIGH, MEDIUM, LOW, CLEAR)\n'
    '  - Date range (from/to)\n'
    '\n'
    'Results pulled from the SQLite database (database/transcripts.db). '
    'Click any result to expand the full transcript, translation, and ISUM.'
)

add_heading(doc, '4.4  [D] DASHBOARD Tab', 2)
doc.add_paragraph(
    'Visual analytics overview of all processed intercepts.\n'
    '\n'
    'Charts and metrics shown:\n'
    '  - Threat level distribution (CRITICAL -> CLEAR breakdown)\n'
    '  - Language distribution (pie or bar chart)\n'
    '  - Processing time trends over time\n'
    '  - Model confidence trends (ASR, LangID, translation)\n'
    '\n'
    'Useful for situational awareness briefings and command-level overviews.'
)

add_heading(doc, '4.5  [H] HISTORY Tab', 2)
doc.add_paragraph(
    'Chronological list of all past intercepts stored in the database.\n'
    '\n'
    'Displayed per entry:\n'
    '  - Report ID and timestamp\n'
    '  - Detected language and threat level\n'
    '  - Short translation summary\n'
    '\n'
    'Allows analysts to quickly browse recent intercepts without re-processing. '
    'Powered by db.get_all_intercepts() database query.'
)

add_heading(doc, '4.6  [E] EXPORT Tab', 2)
doc.add_paragraph(
    'Download processed intercept reports in multiple formats.\n'
    '\n'
    'PDF Export (ReportLab):\n'
    '  - Professional military-format layout\n'
    '  - Includes: threat badge, metadata table, 5W ISUM table, transcript, translation\n'
    '  - Metrics from the METRICS tab (WER, BLEU, chrF, TER) are embedded if available\n'
    '\n'
    'DOCX Export (python-docx):\n'
    '  - Microsoft Word document format\n'
    '\n'
    'JSON Export:\n'
    '  - Raw pipeline result dictionary\n'
    '  - Useful for downstream processing, archiving, or integration with other systems'
)

add_heading(doc, '4.7  [M] METRICS Tab', 2)
doc.add_paragraph(
    'Quality evaluation for ASR and translation outputs.\n'
    '\n'
    'TIER 1 - Automatic (no reference needed, computed automatically):\n'
    '  - RTF (Real-Time Factor): processing speed vs audio duration\n'
    '  - Segment confidence: mean, std deviation, % low/high confidence segments\n'
    '  - Model agreement: 3-way LangID agreement score (Whisper + FastText + MMS)\n'
    '  - Ensemble LangID score\n'
    '  - 5W ISUM completeness score (0-4, one point per answered 5W question)\n'
    '  - Keyword density, stage timings, memory usage (start/peak MB)\n'
    '  - Vocabulary richness (TTR - Type-Token Ratio)\n'
    '  - Back-translation chrF: NLLB translates back to source language and compares to original transcript\n'
    '\n'
    'TIER 2 - Reference-based (analyst provides ground truth text):\n'
    '  - WER (Word Error Rate) and CER (Character Error Rate) via jiwer library\n'
    '  - BLEU, chrF, TER (Translation Error Rate) via sacrebleu library\n'
    '  - Entered manually by the analyst\n'
    '  - Stored in session state and included in PDF/DOCX export when available'
)

add_heading(doc, '4.8  [A] ANNOTATE Tab', 2)
doc.add_paragraph(
    'Human-in-the-loop correction and training data collection.\n'
    '\n'
    'What analysts can correct:\n'
    '  - Transcript (ASR output)\n'
    '  - Translation\n'
    '  - Detected language\n'
    '  - 5W fields (Who, What, Where, When, Assessment)\n'
    '  - Threat level\n'
    '\n'
    'Change tracking:\n'
    'Flags stored per annotation: transcript_changed, translation_changed, isum_changed\n'
    '\n'
    'Training data export:\n'
    'Annotations exported as JSON for fine-tuning with HuggingFace Trainer / Axolotl / TRL\n'
    'Format: {asr: [...], translation: [...], isum: [...]}\n'
    '\n'
    'Stats panel shows:\n'
    '  - Total annotations in database\n'
    '  - Breakdown by type (ASR / Translation / ISUM corrections)\n'
    '  - Breakdown by language'
)

# ================================================================
# SECTION 5: DATABASE
# ================================================================
add_heading(doc, '5. Database Structure', 1)
doc.add_paragraph('File: database/transcripts.db (SQLite - fully local, no server required)\n')

add_table(doc,
    ['Table', 'Key Columns', 'Purpose'],
    [
        ('intercepts', 'id, report_id, audio_file, final_language, transcript, translation, threat_level',
         'One row per processed audio file.'),
        ('segments', 'intercept_id, start_sec, end_sec, text, confidence, no_speech_prob',
         'Word-level / segment-level ASR output with timestamps.'),
        ('keyword_alerts', 'intercept_id, category, severity, matched_word, matched_in, start_sec',
         'Each keyword match stored as a separate row.'),
        ('isums', 'intercept_id, report_id, who, what, where, when_field, assessment, threat_level',
         'ISUM 5W report fields extracted from LLM or rule-based output.'),
        ('annotations', 'intercept_id, corrected_transcript, corrected_translation, corrected_language, isum_changed',
         'Analyst corrections for model fine-tuning.'),
    ]
)

# ================================================================
# SECTION 6: LANGUAGE ROUTING
# ================================================================
add_heading(doc, '6. Language Routing Summary', 1)
doc.add_paragraph('How the system decides which translation model to use for each intercept:\n')

add_table(doc,
    ['Language', 'Routing Set', 'Action'],
    [
        ('English (en)', 'ENGLISH_LIKE', 'No translation. Transcript passed directly to ISUM.'),
        ('Dogri (doi)', 'INDIC_LANGS', 'IndicTrans2 (indictrans2-indic-en-1B). Only language not in NLLB-200.'),
        ('Hindi (hi), Punjabi (pa), Urdu (ur), Pashto (ps), Arabic (ar), Chinese (zh), and 15+ others',
         'NLLB_LANGS',
         'NLLB-200 (nllb-200-distilled-600M). Covers all configured threat-area languages.'),
    ]
)

# ================================================================
# SAVE
# ================================================================
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'VANI_System_Documentation.docx')
doc.save(out_path)
print('Saved:', out_path)
