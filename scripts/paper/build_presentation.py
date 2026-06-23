"""
build_presentation.py
Generates VANI_Presentation.docx — a fully formatted interim project presentation.
Run: python build_presentation.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0D, 0x1B, 0x2A)
GREEN       = RGBColor(0x00, 0x77, 0x44)
GREEN_DARK  = RGBColor(0x00, 0x55, 0x33)
GOLD        = RGBColor(0xB8, 0x86, 0x00)
RED         = RGBColor(0xC6, 0x28, 0x28)
ORANGE      = RGBColor(0xE6, 0x51, 0x00)
BLUE        = RGBColor(0x15, 0x65, 0xC0)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF4, 0xF7, 0xFA)
MID_GREY    = RGBColor(0x5A, 0x70, 0x80)
DARK_GREY   = RGBColor(0x1A, 0x25, 0x35)

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)

def set_cell_border(cell, **borders):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, colour in borders.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "6")
        el.set(qn("w:color"), colour)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=10,
            colour=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    return run

def slide_heading(doc, text, level=1):
    """Navy background heading bar."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "0D1B2A")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    cell.width = Inches(6.5)
    r = p.add_run(f"  {text}")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16 if level == 1 else 13)
    r.font.color.rgb = WHITE
    doc.add_paragraph()

def sub_heading(doc, text, colour=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, bold=True, size=12, colour=colour)
    return p

def body(doc, text, bullet=False, size=10, colour=DARK_GREY, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if bullet:
        p.style = doc.styles["List Bullet"]
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, text, size=size, colour=colour)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), "00AA66")
    pBdr.append(bottom)
    pPr.append(pBdr)

def simple_table(doc, headers, rows, col_widths=None,
                 hdr_bg="0D1B2A", alt_bg="F4F7FA"):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, hdr_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, size=9, colour=WHITE)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = tbl.rows[r_idx + 1]
        bg = alt_bg if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if isinstance(val, tuple):          # (text, bold, colour_hex)
                text, bold, col_hex = val
                rgb = RGBColor(int(col_hex[:2],16),
                               int(col_hex[2:4],16),
                               int(col_hex[4:],16))
                add_run(p, text, bold=bold, size=9, colour=rgb)
            else:
                add_run(p, str(val), size=9, colour=DARK_GREY)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl

def status_badge(p, text, bg_hex, fg_hex):
    """Inline coloured text badge (simulated via bold coloured run)."""
    fg = RGBColor(int(fg_hex[:2],16), int(fg_hex[2:4],16), int(fg_hex[4:],16))
    add_run(p, f"[{text}]", bold=True, size=9, colour=fg)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
c = cover.cell(0, 0)
set_cell_bg(c, "0D1B2A")
c.width = Inches(6.5)

for txt, sz, colour in [
    ("V  A  N  I", 36, WHITE),
    ("Voice Analysis & Neural Intelligence", 16, RGBColor(0xA0,0xB8,0xC8)),
    ("", 8, WHITE),
    ("Military-Grade Offline Radio Intercept Analysis System", 13, RGBColor(0x00,0xCC,0x77)),
    ("", 8, WHITE),
    ("Interim Project Presentation  |  March 2026", 10, RGBColor(0x80,0xA0,0xB0)),
]:
    p = c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, txt, bold=(sz > 20), size=sz, colour=colour, font="Calibri")

doc.add_paragraph()
doc.add_paragraph()

# ── 1. AIM / SCOPE ────────────────────────────────────────────────────────────
slide_heading(doc, "1.  AIM / SCOPE")

sub_heading(doc, "Aim")
body(doc,
     "To develop a fully offline, CPU-deployable intelligence system that automatically "
     "transcribes, translates, and analyses military radio intercepts in multiple languages — "
     "producing structured intelligence summaries (ISUM) in real time, without any dependence "
     "on cloud services or internet connectivity.")

sub_heading(doc, "Problem Statement")
for b in [
    "Manual transcription of radio intercepts is slow, error-prone, and language-constrained.",
    "Existing commercial tools require cloud connectivity — unacceptable in classified field deployments.",
    "Most ASR/NLP systems do not support Indian border-area languages: Punjabi, Dogri, Pashto, Kashmiri.",
    "Raw intercepts provide no structured intelligence — analysts must extract 5W fields manually.",
    "No existing single system combines ASR + Language ID + Translation + Keyword Detection + ISUM offline.",
]:
    body(doc, b, bullet=True)

sub_heading(doc, "Scope")
simple_table(doc,
    ["Dimension", "Coverage"],
    [
        ("Languages", "Hindi · Punjabi · Urdu · Pashto · Nepali · Dogri · Kashmiri · Maithili · Bengali · Sindhi · Sinhala · Mandarin · Burmese · Tibetan · English  (15 languages)"),
        ("Input Formats", "WAV · MP3 · OGG · FLAC · M4A"),
        ("Output", "Structured ISUM (5W), PDF report, DOCX report, JSON"),
        ("Deployment", "Standalone Windows/Linux machine, fully air-gapped"),
        ("Hardware Target", "8 GB RAM, CPU-only — no GPU required"),
        ("Interface", "Web-based GUI (Streamlit), accessible via browser on LAN"),
    ],
    col_widths=[1.8, 4.7]
)

divider(doc)

# ── 2. DELIVERABLES ───────────────────────────────────────────────────────────
slide_heading(doc, "2.  DELIVERABLES")

deliverables = [
    ("D1", "Core Processing Pipeline",
     "End-to-end: VAD → ASR → Language ID → Translation → Keywords → ISUM. 15 languages, fully offline.",
     "✅ COMPLETE", "007744"),
    ("D2", "Web-Based Intelligence Interface",
     "8-tab Streamlit GUI: PROCESS, ISUM REPORT, SEARCH, DASHBOARD, HISTORY, EXPORT, METRICS, ANNOTATE.",
     "✅ COMPLETE", "007744"),
    ("D3", "Structured ISUM & Keyword Detection",
     "804-keyword multilingual dictionary (12 categories). Automated 5W extraction. Threat classification.",
     "✅ COMPLETE", "007744"),
    ("D4", "Quality Metrics & Evaluation Framework",
     "Tier 1 auto metrics (RTF, confidence, model agreement). Tier 2 reference metrics (WER/CER, BLEU/chrF).",
     "✅ COMPLETE", "007744"),
    ("D5", "Training Data Collection System",
     "Analyst annotation interface. Structured export (ASR / translation / ISUM sub-datasets) for fine-tuning.",
     "✅ COMPLETE", "007744"),
    ("D6", "Offline Windows Server Deployment Package",
     "Installation guide, dependency pinning, model download scripts, air-gapped validation.",
     "🔄 IN PROGRESS", "B88600"),
    ("D7", "Fine-Tuned Models on Actual Intercepts",
     "Domain-specific Whisper, NLLB, and Qwen fine-tuning on verified military radio data.",
     "📅 PLANNED", "1565C0"),
]

tbl = doc.add_table(rows=1 + len(deliverables), cols=4)
tbl.style = "Table Grid"
for i, h in enumerate(["ID", "Deliverable", "Description", "Status"]):
    c2 = tbl.rows[0].cells[i]
    set_cell_bg(c2, "0D1B2A")
    p = c2.paragraphs[0]
    add_run(p, h, bold=True, size=9, colour=WHITE)

for r_i, (did, title, desc, status, scol) in enumerate(deliverables):
    bg = "F4F7FA" if r_i % 2 == 0 else "FFFFFF"
    row = tbl.rows[r_i + 1]
    for c_i, (val, bold, sz, col) in enumerate([
        (did,    True,  9, "0D1B2A"),
        (title,  True,  9, "0D1B2A"),
        (desc,   False, 9, "1A2535"),
        (status, True,  9, scol),
    ]):
        cell = row.cells[c_i]
        set_cell_bg(cell, bg)
        rgb = RGBColor(int(col[:2],16), int(col[2:4],16), int(col[4:],16))
        add_run(cell.paragraphs[0], val, bold=bold, size=sz, colour=rgb)

for i, w in enumerate([0.4, 1.4, 3.5, 1.2]):
    for row in tbl.rows:
        row.cells[i].width = Inches(w)
doc.add_paragraph()

divider(doc)

# ── 3. PHASE-WISE IMPLEMENTATION ─────────────────────────────────────────────
slide_heading(doc, "3.  PHASE-WISE IMPLEMENTATION")

phases = [
    ("Phase 1\nFOUNDATION\nNov–Dec 2025",
     "• System architecture design\n• Base pipeline (VAD → ASR → Translate)\n"
     "• Streamlit UI skeleton (3 tabs)\n• SQLite database schema\n• Config-driven model loading",
     "✅ COMPLETE", "007744"),
    ("Phase 2\nLANG ID\nJan 2026",
     "• 3-way Language ID voting system\n• NLLB-200 translation integration\n"
     "• IndicTrans2 patches (transformers 5.x)\n• Punjabi/Hindi disambiguation fix\n• MMS-LID audio-based language vote",
     "✅ COMPLETE", "007744"),
    ("Phase 3\nINTEL FUSION\nFeb 2026",
     "• 804-keyword multilingual dictionary (12 categories)\n• Rule-based ISUM (5W extraction)\n"
     "• Qwen2.5 LLM ISUM (optional GPU mode)\n• Threat level classification\n"
     "• WHO/WHERE/WHEN pattern improvements",
     "✅ COMPLETE", "007744"),
    ("Phase 4\nOUTPUT\nFeb–Mar 2026",
     "• PDF/DOCX professional military report export\n• Metrics framework (RTF, WER, BLEU, chrF)\n"
     "• Analyst annotation tab (training data)\n• Model caching (10× speedup on 2nd run)",
     "✅ COMPLETE", "007744"),
    ("Phase 5\nTRAIN & DEPLOY\nApr–Jun 2026",
     "• Windows Server deployment package\n• Offline fine-tuning on actual military intercepts\n"
     "• Domain-adaptive Whisper fine-tuning\n• Evaluation on classified test set",
     "🔄 IN PROGRESS", "B88600"),
    ("Phase 6\nSCALE\nJul–Sep 2026",
     "• GPU-accelerated deployment\n• Real-time streaming mode\n"
     "• Multi-intercept batch processing\n• Integration with existing SIGINT infrastructure",
     "📅 PLANNED", "1565C0"),
]

tbl2 = doc.add_table(rows=1 + len(phases), cols=3)
tbl2.style = "Table Grid"
for i, h in enumerate(["Phase", "Activities", "Status"]):
    c3 = tbl2.rows[0].cells[i]
    set_cell_bg(c3, "0D1B2A")
    p = c3.paragraphs[0]
    add_run(p, h, bold=True, size=9, colour=WHITE)

for r_i, (phase, acts, status, scol) in enumerate(phases):
    bg = "F4F7FA" if r_i % 2 == 0 else "FFFFFF"
    row = tbl2.rows[r_i + 1]
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    set_cell_bg(row.cells[2], bg)
    add_run(row.cells[0].paragraphs[0], phase, bold=True, size=9,
            colour=RGBColor(0x0D,0x1B,0x2A))
    add_run(row.cells[1].paragraphs[0], acts, size=9,
            colour=RGBColor(0x1A,0x25,0x35))
    rgb_s = RGBColor(int(scol[:2],16), int(scol[2:4],16), int(scol[4:],16))
    add_run(row.cells[2].paragraphs[0], status, bold=True, size=9, colour=rgb_s)

for i, w in enumerate([1.3, 4.4, 1.2]):
    for row in tbl2.rows:
        row.cells[i].width = Inches(w)
doc.add_paragraph()

divider(doc)

# ── 4. WORKFLOW ───────────────────────────────────────────────────────────────
slide_heading(doc, "4.  SYSTEM WORKFLOW")

sub_heading(doc, "End-to-End Pipeline (8 Stages)")
pipeline_steps = [
    ("Stage 1 — VAD (Voice Activity Detection)",
     "Silero VAD 4.0 removes silence and radio static. Speech segments extracted. "
     "Typical: ~0.6s for 5 min audio."),
    ("Stage 2 — Audio Preprocessing",
     "librosa + noisereduce: normalise to 16 kHz mono, noise gate at –20 dB, "
     "stationary noise suppression for radio channel artefacts."),
    ("Stage 3 — VAD-Aware Chunking",
     "Audio split into ≤29s chunks aligned to VAD speech boundaries. "
     "Prevents mid-word cuts. Chunk metadata (start_sec, end_sec) preserved."),
    ("Stage 4 — ASR (Automatic Speech Recognition)",
     "Whisper Large-v3-Turbo (CTranslate2, int8 quantised). beam=2, temperature=0.0. "
     "Initial prompt injects military vocabulary + Punjabi phrases. "
     "Word-level timestamps returned per segment. Language detected from first chunk and reused."),
    ("Stage 5 — Language Identification",
     "3-way confidence-weighted voting: "
     "(a) Whisper language probability, "
     "(b) FastText lid.176.bin on transcript text, "
     "(c) MMS-LID-256 on audio. "
     "Punjabi fix: Gurmukhi script detection OR FastText/MMS=pa + Whisper=hi → forces pa → NLLB route."),
    ("Stage 6 — Translation",
     "English: skip. Dogri: IndicTrans2. All others: NLLB-200-distilled-600M → English. "
     "Back-translation via NLLB (EN → source) for chrF round-trip quality scoring."),
    ("Stage 7 — Keyword Detection",
     "804 keywords, 12 threat categories, 8 languages. "
     "Regex word-boundary matching on both transcript AND translation. "
     "Segment-level timing attached to every alert. Threat level: CRITICAL/HIGH/MEDIUM/LOW/CLEAR."),
    ("Stage 8 — ISUM Generation",
     "Rule-based (default): WHO (callsigns, ranks, unit designators), "
     "WHERE (grids, MGRS, terrain, distance+bearing), "
     "WHEN (military time, H-hour, relative), "
     "ASSESSMENT (category-aware tactical text). "
     "Qwen2.5-1.5B LLM mode available for GPU deployments."),
]

for title, desc in pipeline_steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.2)
    add_run(p, f"▸  {title}", bold=True, size=10, colour=GREEN)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.4)
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, desc, size=9, colour=DARK_GREY)

sub_heading(doc, "Translation Routing Logic")
simple_table(doc,
    ["Language", "Route", "Model", "Reason"],
    [
        ("English (en)", "Skip", "None", "Already in target language"),
        ("Dogri (doi)", "IndicTrans2", "indictrans2-indic-en-1B", "Not in NLLB-200 vocabulary"),
        ("Hindi, Punjabi, Urdu, Nepali,\nBengali, Kashmiri, Maithili,\nSindhi, Sinhala", "NLLB", "nllb-200-distilled-600M", "IndicTrans2 DynamicCache incompatibility with transformers 5.x"),
        ("Pashto, Mandarin, Burmese,\nTibetan, Persian, Arabic", "NLLB", "nllb-200-distilled-600M", "Primary NLLB-supported languages"),
    ],
    col_widths=[1.5, 1.0, 1.8, 2.2]
)

divider(doc)

# ── 5. CURRENT STATUS ─────────────────────────────────────────────────────────
slide_heading(doc, "5.  CURRENT STATUS")

sub_heading(doc, "Pipeline Performance (CPU, 8 GB RAM)")
simple_table(doc,
    ["Stage", "Technology", "Time (5s audio)", "Notes"],
    [
        ("VAD",          "Silero VAD 4.0",                       "~0.6s",   "Speech/silence separation"),
        ("Preprocessing","librosa + noisereduce",                 "~0.3s",   "16kHz, noise reduction"),
        ("Chunking",     "VAD-aware splitter",                   "~0.2s",   "Max 29s chunks"),
        ("ASR",          "Whisper Large-v3-Turbo (CTranslate2)", "~40s (1st run)\n~5s (cached)", "int8 quantised, beam=2"),
        ("Language ID",  "FastText + MMS-LID (cached)",         "~0.5s",   "3-way vote"),
        ("Translation",  "NLLB-200-distilled-600M",              "~3–5s",   "Load-on-demand"),
        ("Keywords",     "Regex (804 entries)",                  "< 0.1s",  "12 categories"),
        ("ISUM",         "Rule-based",                           "< 0.1s",  "Qwen available for GPU"),
        ("TOTAL",        "Full pipeline",                        "~15–25s (2nd run)", "Model caching active"),
    ],
    col_widths=[1.2, 2.2, 1.6, 1.5]
)

sub_heading(doc, "Models Installed (All Offline)")
simple_table(doc,
    ["Model", "Size", "Purpose", "Status"],
    [
        ("Whisper Large-v3-Turbo (CTranslate2)", "~1.6 GB", "ASR",                        ("✅ Operational", True, "007744")),
        ("NLLB-200-distilled-600M",              "~2.2 GB", "Translation (all languages)", ("✅ Operational", True, "007744")),
        ("IndicTrans2-indic-en-1B",              "~3.5 GB", "Translation (Dogri)",         ("✅ Operational", True, "007744")),
        ("FastText lid.176.bin",                 "~900 MB", "Text Language ID",            ("✅ Operational", True, "007744")),
        ("MMS-LID-256",                          "~150 MB", "Audio Language ID",           ("✅ Operational", True, "007744")),
        ("Qwen2.5-1.5B-Instruct",               "~3.1 GB", "LLM ISUM (GPU mode)",         ("⚡ GPU Needed",  True, "B88600")),
        ("TOTAL",                                "~11.5 GB","",                            ("", False, "000000")),
    ],
    col_widths=[2.2, 0.9, 1.8, 1.6]
)

sub_heading(doc, "Web Interface — 8 Tabs")
tabs = [
    ("[P] PROCESS",     "Upload audio, run pipeline, real-time stage-by-stage progress bar"),
    ("[I] ISUM REPORT", "5W intelligence summary, threat badge, confidence flags, PDF download"),
    ("[S] SEARCH",      "Fuzzy full-text search with language / threat / date-range filters"),
    ("[D] DASHBOARD",   "Threat distribution charts, language breakdown, recent priority intercepts"),
    ("[H] HISTORY",     "Browse all processed intercepts, sortable by threat / language / date"),
    ("[X] EXPORT",      "Single & batch PDF / DOCX / JSON export with metrics included"),
    ("[M] METRICS",     "RTF, ASR confidence histogram, model agreement, WER/CER, BLEU/chrF"),
    ("[A] ANNOTATE",    "Analyst correction of transcript / translation / ISUM for training data"),
]
for tab, desc in tabs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.2)
    add_run(p, f"{tab}:  ", bold=True, size=10, colour=NAVY)
    add_run(p, desc, size=10, colour=DARK_GREY)

divider(doc)

# ── 6. KEY FEATURES ───────────────────────────────────────────────────────────
slide_heading(doc, "6.  KEY FEATURES")

features = [
    ("Fully Air-Gapped Operation",
     "All models run locally. Zero internet dependency enforced in code (HF_HUB_OFFLINE=1). "
     "Suitable for classified field deployments and secure server environments."),
    ("3-Way Language Identification Voting",
     "Whisper (ASR) + FastText (text, 176 languages) + MMS-LID (audio, 256 languages). "
     "Confidence-weighted voting with Unanimous / Majority / Best-confidence fallback. "
     "Every result includes a vote_note explaining the decision."),
    ("Punjabi / Hindi Disambiguation",
     "Whisper frequently misidentifies Punjabi audio as Hindi. "
     "Corrected via: (a) Gurmukhi Unicode block detection in transcript, "
     "(b) FastText/MMS both returning pa when Whisper says hi. "
     "Forces pa → NLLB translation route."),
    ("Tactical WHO / WHERE / WHEN Extraction",
     "WHO: Full NATO phonetic alphabet, numeric callsigns (TF-7, SF-3), "
     "military ranks (Captain, Subedar, Havildar), unit formations. "
     "WHERE: MGRS grid refs, bare numeric grids, distance+bearing ('5 km north'), "
     "terrain features (ridgeline, bunker, outpost). "
     "WHEN: Military time (0530 hrs), H-hour notation (H+30), "
     "relative time ('in 30 minutes'), day/date patterns."),
    ("Category-Aware Intelligence Assessment",
     "Assessment text is specific to the top threat categories detected — not boilerplate. "
     "Examples: CBRN → 'escalate to higher command immediately'; "
     "Attack → 'active fire mission underway — immediate action required'; "
     "Support request → 'assess force disposition'."),
    ("804-Keyword Multilingual Threat Detection",
     "12 categories: enemy_activity (CRITICAL), attack (CRITICAL), explosives (CRITICAL), "
     "nuclear_chem_bio (CRITICAL), weapons (HIGH), movement (HIGH), location (HIGH), "
     "support_request (HIGH), pre_attack (HIGH), casualties (HIGH), command (MEDIUM), comms (LOW). "
     "Searches both original transcript AND English translation simultaneously."),
    ("Quality Metrics Framework",
     "Tier 1 (auto): RTF, ASR confidence distribution, 3-way model agreement, "
     "5W completeness score, stage timings, memory usage, vocabulary richness (TTR), "
     "back-translation chrF. "
     "Tier 2 (analyst reference): WER/CER (jiwer), BLEU/chrF/TER (sacrebleu)."),
    ("Training Data Collection Pipeline",
     "Annotation tab collects analyst-corrected transcripts, translations, and ISUM fields. "
     "Exported as structured JSON: asr[], translation[], isum[] sub-datasets. "
     "Compatible with HuggingFace Trainer, Axolotl, and TRL for fine-tuning."),
    ("Model Caching for Speed",
     "Whisper, FastText, and MMS-LID cached via @st.cache_resource on first load. "
     "Subsequent pipeline runs skip model loading — reduces processing time from "
     "~65s to ~15–25s for a 5-second audio file."),
]

for i, (title, desc) in enumerate(features):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    add_run(p, f"{'0' if i<9 else ''}{i+1}.  {title}", bold=True, size=11, colour=GREEN)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(3)
    add_run(p2, desc, size=9, colour=DARK_GREY)

divider(doc)

# ── 7. LITERATURE REVIEW ──────────────────────────────────────────────────────
slide_heading(doc, "7.  LITERATURE REVIEW — REFERENCE PAPERS")

papers = [
    {
        "num": "Paper 1",
        "title": "Robust Speech Recognition via Large-Scale Weak Supervision",
        "authors": "Alec Radford, Jong Wook Kim, Tao Xu, Greg Brockman, Christine McLeavey, Ilya Sutskever",
        "venue": "OpenAI Technical Report  |  arXiv:2212.04356  |  2022",
        "module": "ASR Module — Whisper Large-v3-Turbo",
        "relevance": (
            "Whisper's encoder-decoder transformer (trained on 680,000 hours of multilingual audio) "
            "provides the ASR backbone for VANI. The paper establishes zero-shot multilingual "
            "transcription without fine-tuning — critical for languages with no labelled training data. "
            "Key techniques adopted: initial prompt injection for military vocabulary, "
            "condition_on_previous_text=False to prevent hallucination loops on radio static, "
            "and word-level timestamps for precision segment localisation in search."
        ),
    },
    {
        "num": "Paper 2",
        "title": "No Language Left Behind: Scaling Human-Centered Machine Translation",
        "authors": "Marta R. Costa-jussà, James Cross, Onur Çelebi et al.  (Meta AI Research)",
        "venue": "arXiv:2207.04672  |  2022",
        "module": "Translation Module — NLLB-200-distilled-600M",
        "relevance": (
            "NLLB-200 supports 200 languages in a single model — including all Indic languages "
            "in VANI (Hindi, Punjabi, Urdu, Nepali, Bengali, Kashmiri, Sindhi, Sinhala). "
            "The distilled 600M variant fits within the 8 GB RAM constraint. "
            "The paper's findings on low-resource language translation quality directly motivated "
            "routing all Indic languages through NLLB after IndicTrans2's DynamicCache "
            "incompatibility with transformers 5.x was discovered. "
            "VANI also uses NLLB for back-translation chrF scoring (EN → source → compare to original)."
        ),
    },
    {
        "num": "Paper 3",
        "title": "IndicTrans2: Towards High-Quality and Accessible Machine Translation Models for all 22 Scheduled Indian Languages",
        "authors": "Jay Gala, Pranjal A. Chitale, A K Raghavan, Varun Gumma, Sumanth Doddapaneni et al.  (AI4Bharat)",
        "venue": "Transactions on Machine Learning Research (TMLR)  |  2023",
        "module": "Translation Module — IndicTrans2-indic-en-1B (Dogri fallback)",
        "relevance": (
            "IndicTrans2 provides state-of-the-art translation quality for all 22 Indian scheduled "
            "languages, including Dogri — a language not covered by NLLB-200. "
            "VANI uses IndicTrans2 as the dedicated route for Dogri (doi). "
            "The paper's custom tokeniser with Gurmukhi/Devanagari/Arabic script support "
            "informed VANI's Unicode-block script detection used in Punjabi disambiguation. "
            "Multiple compatibility patches were applied to run IndicTrans2 under transformers>=5.0: "
            "onnx stub, tie_weights() **kwargs, use_cache=False, attn_implementation=eager."
        ),
    },
    {
        "num": "Paper 4",
        "title": "Scaling Speech Technology to 1,000+ Languages",
        "authors": "Vineel Pratap, Andros Tjandra, Bowen Shi, Paden Tomasello et al.  (Meta AI Research)",
        "venue": "arXiv:2305.13516  |  2023",
        "module": "Language ID Module — MMS-LID-256 (audio-based vote)",
        "relevance": (
            "The Massively Multilingual Speech (MMS) project produced MMS-LID-256, "
            "a language identification model trained on 256 languages from audio alone — "
            "without requiring any text transcript. "
            "VANI incorporates this as the third vote in its language identification ensemble, "
            "specifically to resolve cases where text-based FastText is confused by "
            "romanised Indic script (e.g. Punjabi transcribed in Latin characters by Whisper). "
            "The audio-based vote proved most reliable for languages with non-Latin scripts "
            "and was critical for improving 3-way voting accuracy for Punjabi and Pashto."
        ),
    },
    {
        "num": "Paper 5",
        "title": "Bag of Tricks for Efficient Text Classification",
        "authors": "Armand Joulin, Edouard Grave, Piotr Bojanowski, Tomas Mikolov  (Facebook AI Research)",
        "venue": "Proceedings of EACL 2017  |  arXiv:1607.01759",
        "module": "Language ID Module — FastText lid.176.bin (text-based vote)",
        "relevance": (
            "FastText's sub-word character n-gram text classification forms the second vote "
            "in VANI's language identification system. The lid.176.bin model (176 languages) "
            "operates on the Whisper-produced transcript and returns ranked language probabilities. "
            "The paper's key insight — character n-gram features capture morphological patterns "
            "across language families — makes FastText effective for Indic languages which share "
            "substantial morphological overlap (Hindi-Urdu, Hindi-Dogri) while differing in script. "
            "FastText is the fastest of the three voting models (<1ms per call) making it "
            "the ideal second vote alongside the heavier Whisper and MMS models."
        ),
    },
]

for paper in papers:
    # Coloured section bar
    tbl_p = doc.add_table(rows=1, cols=1)
    tbl_p.alignment = WD_TABLE_ALIGNMENT.LEFT
    cp = tbl_p.cell(0, 0)
    set_cell_bg(cp, "003322")
    cp.width = Inches(6.5)
    pp = cp.paragraphs[0]
    add_run(pp, f"  {paper['num']} — {paper['title']}", bold=True, size=10, colour=RGBColor(0x00,0xDD,0x88))
    pp.paragraph_format.space_before = Pt(2)
    pp.paragraph_format.space_after  = Pt(2)

    p_auth = doc.add_paragraph()
    p_auth.paragraph_format.left_indent = Inches(0.2)
    p_auth.paragraph_format.space_before = Pt(2)
    add_run(p_auth, paper["authors"], italic=True, size=9, colour=MID_GREY)

    p_venue = doc.add_paragraph()
    p_venue.paragraph_format.left_indent = Inches(0.2)
    add_run(p_venue, paper["venue"], size=9, colour=BLUE)

    p_mod = doc.add_paragraph()
    p_mod.paragraph_format.left_indent = Inches(0.2)
    add_run(p_mod, "Used in: ", bold=True, size=9, colour=GREEN)
    add_run(p_mod, paper["module"], size=9, colour=DARK_GREY)

    p_rel = doc.add_paragraph()
    p_rel.paragraph_format.left_indent = Inches(0.2)
    p_rel.paragraph_format.space_after = Pt(8)
    add_run(p_rel, "Relevance: ", bold=True, size=9, colour=NAVY)
    add_run(p_rel, paper["relevance"], size=9, colour=DARK_GREY)

divider(doc)

# ── 8. PAPER PUBLICATION DETAILS ─────────────────────────────────────────────
slide_heading(doc, "8.  PAPER PUBLICATION DETAILS")

sub_heading(doc, "Planned Paper Title")
p_title = doc.add_paragraph()
p_title.paragraph_format.left_indent = Inches(0.2)
add_run(p_title,
        '"VANI: An Offline Multilingual Spoken Intelligence Analysis System for '
        'Low-Resource Indian Border Languages"',
        bold=True, size=11, colour=NAVY)

sub_heading(doc, "Key Contributions")
contribs = [
    "3-way confidence-weighted Language ID voting combining ASR, text, and audio models with Punjabi script disambiguation",
    "Rule-based military ISUM generation with domain-specific tactical pattern extraction (WHO/WHERE/WHEN)",
    "Indic-language translation routing solving IndicTrans2 / NLLB-200 complementarity for 15 languages",
    "End-to-end offline deployment of 11.5 GB model stack on 8 GB RAM CPU-only hardware",
    "Military annotation framework for building labelled intercept training data (ASR + translation + ISUM)",
]
for c in contribs:
    body(doc, c, bullet=True)

sub_heading(doc, "Target Publication Venues")
simple_table(doc,
    ["Priority", "Venue", "Scope", "Impact Factor"],
    [
        ("Primary",     "IEEE Transactions on Information Forensics and Security",
         "Intelligence systems, signal processing, security",         "~7.2"),
        ("Alternative", "INTERSPEECH 2026 / ICASSP 2026",
         "Multilingual ASR, low-resource speech processing",          "Top-tier conference"),
        ("Alternative", "ACL 2026 / EMNLP 2026",
         "Multilingual NLP, annotation frameworks, machine translation","Top-tier conference"),
    ],
    col_widths=[0.8, 2.5, 2.2, 1.0]
)

divider(doc)

# ── 9. HARDWARE & SOFTWARE PROCUREMENT ───────────────────────────────────────
slide_heading(doc, "9.  PROGRESS ON HARDWARE & SOFTWARE PROCUREMENT")

sub_heading(doc, "Current Development Machine")
simple_table(doc,
    ["Component", "Specification", "Status"],
    [
        ("CPU",     "x86-64 multi-core",                             "✅ Available"),
        ("RAM",     "8 GB",                                          "✅ Sufficient for CPU inference"),
        ("Storage", "~50 GB free (models: ~11.5 GB)",                "✅ Available"),
        ("GPU",     "None",                                          "❌ CPU-only mode"),
        ("OS",      "Windows 11",                                    "✅ Operational"),
        ("Network", "LAN access to localhost:8501",                  "✅ Accessible"),
    ],
    col_widths=[1.3, 3.2, 2.0]
)

sub_heading(doc, "Target Windows Server (Deployment)")
simple_table(doc,
    ["Component", "Minimum", "Recommended", "Reason"],
    [
        ("CPU",      "8-core x86-64",      "16-core",            "Parallel model inference"),
        ("RAM",      "16 GB",              "32 GB",              "Multiple models cached simultaneously"),
        ("Storage",  "100 GB SSD",         "500 GB NVMe",        "Models + audio archive + database"),
        ("GPU",      "None (CPU mode)",    "NVIDIA 16 GB VRAM",  "Enable Qwen ISUM + GPU-accelerated ASR"),
        ("OS",       "Windows Server 2019","Windows Server 2022","Standalone offline deployment"),
        ("Network",  "Air-gapped LAN",     "Air-gapped LAN",     "No internet connectivity required"),
    ],
    col_widths=[1.1, 1.3, 1.5, 2.6]
)

sub_heading(doc, "Software Stack (All Open-Source, Zero Licensing Cost)")
simple_table(doc,
    ["Category", "Package", "Version", "Purpose"],
    [
        ("Deep Learning",   "PyTorch (CPU wheels)",          "2.2.2",  "Neural network inference"),
        ("ASR",             "faster-whisper + CTranslate2",  "Latest", "Speech-to-text"),
        ("NLP / Translation","HuggingFace Transformers",     "4.40.0", "NLLB, IndicTrans2, Qwen"),
        ("Language ID",     "fasttext-wheel",                "0.9.2",  "Pre-built wheel, no compiler"),
        ("UI Framework",    "Streamlit",                     "1.34.0", "Web interface"),
        ("Database",        "SQLite3",                       "Native", "Intercept history"),
        ("PDF Export",      "ReportLab",                     "4.0.7",  "Military-format PDF"),
        ("DOCX Export",     "python-docx",                   "Latest", "Word document export"),
        ("Audio",           "librosa + soundfile + noisereduce","Pinned","Audio processing"),
        ("Metrics",         "jiwer + sacrebleu",             "Pinned", "WER, BLEU, chrF"),
        ("Protobuf",        "protobuf",                      "3.20.3", "Pinned — IndicTrans2 requires ≤3.x"),
    ],
    col_widths=[1.3, 2.0, 0.8, 2.4]
)

sub_heading(doc, "Procurement Status Summary")
for item in [
    "✅  All software: open-source, downloaded, operational — zero licensing cost",
    "✅  All AI models: downloaded and running offline (~11.5 GB total)",
    "🔄  Windows Server hardware: procurement in progress",
    "📅  GPU card (NVIDIA 16 GB+): to be requested — required for Phase 5 training",
    "📅  Secure storage for classified intercept audio: to be provisioned",
]:
    body(doc, item, bullet=False, size=10)

divider(doc)

# ── 10. FUTURE IMPROVEMENTS ───────────────────────────────────────────────────
slide_heading(doc, "10. FUTURE IMPROVEMENTS")

sub_heading(doc, "A. Offline Fine-Tuning on Actual Military Intercepts  (Phase 5)")
body(doc,
     "The annotation system (ANNOTATE tab) is already collecting analyst-corrected data. "
     "Once 1,000+ samples per language are collected, fine-tuning will be executed offline "
     "on the Windows Server.", size=10)
simple_table(doc,
    ["Model", "Framework", "Training Data Source", "Target Metric"],
    [
        ("Whisper Large-v3-Turbo","HuggingFace Seq2SeqTrainer","asr[] from annotation export","WER < 10% on military vocabulary"),
        ("NLLB-200-600M",         "HuggingFace Seq2SeqTrainer","translation[] from annotation","chrF > 55 for Punjabi/Urdu/Pashto"),
        ("Qwen2.5-1.5B",          "TRL SFTTrainer / Axolotl",  "isum[] from annotation export","5W completeness > 90%"),
    ],
    col_widths=[1.4, 1.8, 2.1, 1.2]
)

sub_heading(doc, "B. Short-Term (0–3 months)")
for item in [
    "Real-Time Streaming Mode — process audio chunk-by-chunk as it arrives; sub-10s latency for live monitoring",
    "GPU-Accelerated Deployment — Whisper on GPU: ~5× speedup; Qwen ISUM: from 10 min → 30 seconds",
    "Auto-Escalation Alerts — if CRITICAL + confidence > 0.90, auto-flag for immediate analyst review",
    "Expanded Language Coverage — Balochi, Tibetan dialects, local Punjabi dialect variants (Majhi, Malwi)",
]:
    body(doc, item, bullet=True)

sub_heading(doc, "C. Medium-Term (3–6 months)")
for item in [
    "Speaker Diarisation — identify number of speakers, assign dialogue turns (Speaker A / Speaker B)",
    "Voice Biometrics — build voiceprint library, identify known speakers across multiple intercepts",
    "Confidence-Based Priority Queue — automatically rank intercepts for analyst review by threat × confidence",
    "Multi-Intercept Timeline Reconstruction — correlate intercepts by callsign/location over time",
]:
    body(doc, item, bullet=True)

sub_heading(doc, "D. Long-Term (6–12 months)")
for item in [
    "Multi-Intercept Correlation — network graph of communication participants, cross-reference entities",
    "LLM Contextual ISUM — Qwen fine-tuned on verified intercepts; feed multiple intercepts simultaneously for operational-level picture",
    "GIS Integration — map location coordinates extracted from ISUM onto geospatial display",
    "STANAG-Compatible Output — standardised NATO message format for integration with existing SIGINT infrastructure",
]:
    body(doc, item, bullet=True)

divider(doc)

# ── 11. MISC ASPECTS ─────────────────────────────────────────────────────────
slide_heading(doc, "11. MISC ASPECTS")

sub_heading(doc, "Security Considerations")
simple_table(doc,
    ["Aspect", "Implementation"],
    [
        ("Air-gap enforcement",  "HF_HUB_OFFLINE=1, TRANSFORMERS_OFFLINE=1 set in code — no outbound connections"),
        ("Local processing",     "All models run on-device; no audio or transcript data leaves the machine"),
        ("Database",             "SQLite with WAL mode; no network database; file-level access control"),
        ("Audio retention",      "Analyst controls input_audio/ directory; can be wiped after processing"),
        ("Access control",       "Streamlit restricted to localhost or defined LAN IP range"),
        ("Model integrity",      "All models downloaded once from verified HuggingFace sources before air-gap"),
    ],
    col_widths=[1.8, 4.7]
)

sub_heading(doc, "Known Limitations")
for item in [
    "CPU inference is slow — First run ~65s for 5s audio (model loading). Cached runs: ~15–25s. GPU eliminates this.",
    "Rule-based ISUM is pattern-matching — will miss complex, unusual, or code-word phrasing.",
    "Translation quality — lower for Dogri and Kashmiri (low-resource, limited training data).",
    "No fine-tuning on actual military vocabulary yet — models trained on general-domain data.",
    "No speaker diarisation — multi-speaker intercepts produce merged single-speaker transcript.",
    "Qwen LLM ISUM requires GPU — CPU generation takes 5–10 minutes; disabled by default.",
]:
    body(doc, item, bullet=True)

sub_heading(doc, "Evaluation Metrics Summary")
simple_table(doc,
    ["Metric", "Meaning", "Current Target"],
    [
        ("WER",             "Word Error Rate (ASR accuracy)",                       "< 20% (general domain)"),
        ("CER",             "Character Error Rate",                                 "< 12%"),
        ("BLEU",            "Translation quality (n-gram precision)",               "> 25"),
        ("chrF",            "Character F-score for translation",                    "> 50"),
        ("RTF",             "Real-Time Factor (processing / speech duration)",      "< 5× (CPU), < 0.5× (GPU)"),
        ("5W Completeness", "% of ISUM fields (WHO/WHAT/WHERE/WHEN) populated",    "> 75%"),
        ("Model Agreement", "3-way Language ID ensemble score",                     "> 0.70"),
        ("Threat Accuracy", "Correct threat level classification",                  "> 90% on labelled test set"),
    ],
    col_widths=[1.4, 2.7, 2.4]
)

sub_heading(doc, "Operational Deployment Workflow")
steps = [
    ("Field Operator",  "1. Record radio intercept (WAV/MP3)"),
    ("Field Operator",  "2. Upload to VANI via browser (localhost:8501)"),
    ("VANI System",     "3. Pipeline runs automatically — VAD → ASR → LangID → Translate → Keywords → ISUM"),
    ("VANI System",     "4. ISUM Report generated with threat badge and 5W fields"),
    ("Analyst (HQ)",    "5. Review ISUM Report; download PDF for record"),
    ("Analyst (HQ)",    "6. If CRITICAL: escalate to higher command immediately"),
    ("Analyst (HQ)",    "7. Annotate corrections in ANNOTATE tab for training data"),
    ("Data Pipeline",   "8. Periodic export of training data → fine-tuning → improved models"),
]
for actor, step in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, f"[{actor}]  ", bold=True, size=9, colour=GREEN)
    add_run(p, step, size=9, colour=DARK_GREY)

divider(doc)

# ── 12. ACHIEVEMENT SUMMARY ───────────────────────────────────────────────────
slide_heading(doc, "12. WHAT HAS BEEN ACHIEVED — SUMMARY")

achieved = [
    "10-stage end-to-end offline pipeline — fully operational",
    "15 languages supported including all major Indian border-area languages",
    "3-way Language ID voting (Whisper + FastText + MMS-LID) with confidence weighting",
    "Punjabi/Hindi disambiguation via Gurmukhi script detection and model voting",
    "NLLB-200 + IndicTrans2 dual translation routing — IndicTrans2 5.x patches applied",
    "804-keyword multilingual threat detection across 12 categories",
    "Rule-based ISUM with tactical WHO / WHERE / WHEN / ASSESSMENT patterns",
    "Qwen2.5 LLM ISUM integration for GPU deployments",
    "8-tab professional web interface (Streamlit)",
    "Professional PDF + DOCX military-format report export",
    "Comprehensive metrics framework: Tier 1 (auto) + Tier 2 (reference-based)",
    "Analyst annotation system for training data collection (ASR / translation / ISUM)",
    "SQLite persistent database with fuzzy full-text search and date filters",
    "Model caching — 10× speedup on subsequent pipeline runs",
    "All models installed and operational offline (~11.5 GB, CPU-only)",
]

inprogress = [
    "Windows Server deployment package",
    "Offline model fine-tuning on actual military intercepts",
]

planned = [
    "Real-time streaming mode for live intercept monitoring",
    "GPU-accelerated deployment (pending hardware procurement)",
    "Speaker diarisation for multi-party intercepts",
    "Multi-intercept correlation and timeline reconstruction",
]

tbl_sum = doc.add_table(rows=1 + len(achieved) + len(inprogress) + len(planned), cols=2)
tbl_sum.style = "Table Grid"

# header
set_cell_bg(tbl_sum.rows[0].cells[0], "0D1B2A")
set_cell_bg(tbl_sum.rows[0].cells[1], "0D1B2A")
add_run(tbl_sum.rows[0].cells[0].paragraphs[0], "Status", bold=True, size=9, colour=WHITE)
add_run(tbl_sum.rows[0].cells[1].paragraphs[0], "Milestone", bold=True, size=9, colour=WHITE)

for i, (status, col_hex, items) in enumerate([
    ("✅ COMPLETE",     "007744", achieved),
    ("🔄 IN PROGRESS",  "B88600", inprogress),
    ("📅 PLANNED",      "1565C0", planned),
]):
    for j, item in enumerate(items):
        row_idx = 1 + (0 if i==0 else len(achieved) if i==1 else len(achieved)+len(inprogress)) + j
        row = tbl_sum.rows[row_idx]
        bg = "E8F5EE" if i==0 else ("FFFDE7" if i==1 else "E3F2FD")
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        rgb = RGBColor(int(col_hex[:2],16), int(col_hex[2:4],16), int(col_hex[4:],16))
        add_run(row.cells[0].paragraphs[0], status, bold=True, size=9, colour=rgb)
        add_run(row.cells[1].paragraphs[0], item,   size=9,   colour=DARK_GREY)

for col_i, w in enumerate([1.3, 5.2]):
    for row in tbl_sum.rows:
        row.cells[col_i].width = Inches(w)

doc.add_paragraph()

# ── FOOTER ────────────────────────────────────────────────────────────────────
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.space_before = Pt(12)
add_run(fp, "VANI — Voice Analysis & Neural Intelligence  |  Interim Presentation  |  March 2026  |  INTERNAL USE",
        size=8, colour=MID_GREY, italic=True)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = "VANI_Presentation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
