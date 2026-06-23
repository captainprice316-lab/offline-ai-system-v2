"""
build_paper_exports.py
Generates VANI_Paper.docx and VANI_Paper.pdf from VANI_Paper_Revised.md
Run: venv/bin/python build_paper_exports.py
"""

import re
from pathlib import Path

ROOT     = Path(__file__).resolve().parent
SRC      = ROOT / "output" / "VANI_Paper_Revised.md"   # place source .md here
DOCX_OUT = ROOT / "output" / "VANI_Paper.docx"
PDF_OUT  = ROOT / "output" / "VANI_Paper.pdf"

# ── 1. Parse the markdown into a structured list ─────────────────────────────

def parse_md(path):
    """
    Returns list of (type, content) tuples:
      ('h1', str), ('h2', str), ('h3', str), ('h4', str)
      ('para', str)
      ('table', [header_row, ...rows])
      ('bullet', str)
      ('enum', str)
      ('code', str)
      ('hr', None)
      ('footnote', str)   — lines starting with †/‡
    """
    blocks = []
    lines  = path.read_text(encoding="utf-8").splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # heading
        if line.startswith("#### "):
            blocks.append(("h4", line[5:].strip()))
            i += 1; continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1; continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1; continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1; continue

        # table (look-ahead)
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # remove separator rows (---|---)
            rows = [l for l in table_lines
                    if not re.match(r"^\|[\s\-:]+\|", l)]
            parsed = []
            for row in rows:
                cells = [c.strip() for c in row.strip("|").split("|")]
                parsed.append(cells)
            if parsed:
                blocks.append(("table", parsed))
            continue

        # hr
        if re.match(r"^-{3,}$", line.strip()):
            blocks.append(("hr", None))
            i += 1; continue

        # bullet
        if re.match(r"^[-*] ", line):
            blocks.append(("bullet", line[2:].strip()))
            i += 1; continue

        # numbered list
        if re.match(r"^\d+\. ", line):
            blocks.append(("enum", re.sub(r"^\d+\. ", "", line).strip()))
            i += 1; continue

        # footnote lines (start with † or ‡)
        if line.strip().startswith(("†", "‡", "*All VANI")):
            blocks.append(("footnote", line.strip()))
            i += 1; continue

        # blank
        if not line.strip():
            i += 1; continue

        # paragraph
        blocks.append(("para", line.strip()))
        i += 1

    return blocks


def clean(text):
    """Strip markdown bold/italic/code/links for plain text insertion."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*",     r"\1", text)
    text = re.sub(r"`(.+?)`",       r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    # citation markers like [1], [Radford et al., 2022] etc.
    text = re.sub(r"\[\d+\]",       "", text)
    text = re.sub(r"\[Radford.*?\]", "", text)
    text = re.sub(r"\[Costa.*?\]",   "", text)
    text = re.sub(r"\[Pratap.*?\]",  "", text)
    text = re.sub(r"\[Joulin.*?\]",  "", text)
    text = re.sub(r"\[AI4.*?\]",     "", text)
    text = re.sub(r"\[MMS.*?\]",     "", text)
    return text.strip()


# ── 2. Build DOCX ─────────────────────────────────────────────────────────────

def build_docx(blocks):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── helper: set paragraph font ────────────────────────────────────────────
    def set_para_fmt(para, size=11, bold=False, italic=False,
                     color=None, align=None, space_before=0, space_after=6):
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        if align:
            para.alignment = align
        for run in para.runs:
            run.font.size   = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = color

    # ── styles ────────────────────────────────────────────────────────────────
    NAVY = RGBColor(0x1F, 0x39, 0x64)
    DARK = RGBColor(0x26, 0x26, 0x26)

    # ── title block ───────────────────────────────────────────────────────────
    title_text = None
    body_blocks = []
    for typ, content in blocks:
        if typ == "h1" and title_text is None:
            title_text = content
        else:
            body_blocks.append((typ, content))

    if title_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(clean(title_text))
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = NAVY

    # subtitle / author block
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("VANI System — IEEE Conference Paper")
    r.font.size  = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_after = Pt(12)
    r2 = author.add_run("[Author Name] · [Institution] · [email@domain]")
    r2.font.size  = Pt(10)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── body ─────────────────────────────────────────────────────────────────
    in_enum = False
    enum_counter = 0

    for typ, content in body_blocks:
        if typ == "hr":
            # thin horizontal rule via border
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        if typ == "h2":
            p = doc.add_heading("", level=1)
            p.clear()
            run = p.add_run(clean(content))
            run.font.size  = Pt(13)
            run.font.bold  = True
            run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(4)
            in_enum = False
            continue

        if typ == "h3":
            p = doc.add_heading("", level=2)
            p.clear()
            run = p.add_run(clean(content))
            run.font.size  = Pt(11.5)
            run.font.bold  = True
            run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            in_enum = False
            continue

        if typ == "h4":
            p = doc.add_heading("", level=3)
            p.clear()
            run = p.add_run(clean(content))
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.italic = True
            run.font.color.rgb = DARK
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            continue

        if typ == "para":
            p = doc.add_paragraph()
            run = p.add_run(clean(content))
            run.font.size = Pt(11)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            continue

        if typ == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean(content))
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(3)
            in_enum = False
            continue

        if typ == "enum":
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(clean(content))
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(3)
            continue

        if typ == "footnote":
            p = doc.add_paragraph()
            run = p.add_run(clean(content))
            run.font.size   = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_after = Pt(2)
            continue

        if typ == "table":
            rows = content
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = "Table Grid"
            # header row
            for j, cell_text in enumerate(rows[0]):
                cell = table.cell(0, j)
                cell.text = clean(cell_text)
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(clean(cell_text))
                run.font.bold = True
                run.font.size = Pt(9.5)
                # header bg (navy)
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "1F3964")
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            # data rows
            for i, row in enumerate(rows[1:], 1):
                for j in range(num_cols):
                    cell_text = row[j] if j < len(row) else ""
                    cell = table.cell(i, j)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(clean(cell_text))
                    run.font.size = Pt(9.5)
                    if i % 2 == 0:
                        tc   = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:val"),   "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"),  "EEF2FA")
                        tcPr.append(shd)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

    doc.save(str(DOCX_OUT))
    print(f"[DOCX] Saved: {DOCX_OUT}")


# ── 3. Build PDF ──────────────────────────────────────────────────────────────

def build_pdf(blocks):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable,
                                    KeepTogether)
    from reportlab.platypus.flowables import HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # colours
    NAVY   = colors.HexColor("#1F3964")
    LIGHT  = colors.HexColor("#EEF2FA")
    WHITE  = colors.white
    DARK   = colors.HexColor("#262626")
    GREY   = colors.HexColor("#666666")
    RULE   = colors.HexColor("#AAAAAA")

    PAGE_W, PAGE_H = A4
    MARGIN = 2.2 * cm
    COL_W  = PAGE_W - 2 * MARGIN

    styles = getSampleStyleSheet()

    def S(name, parent="Normal", **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    title_style = S("VTitle",
        fontSize=18, leading=22, textColor=NAVY,
        fontName="Helvetica-Bold", alignment=TA_CENTER,
        spaceAfter=4)

    sub_style = S("VSub",
        fontSize=10, leading=13, textColor=GREY,
        fontName="Helvetica-Oblique", alignment=TA_CENTER,
        spaceAfter=2)

    author_style = S("VAuthor",
        fontSize=9.5, leading=12, textColor=GREY,
        fontName="Helvetica", alignment=TA_CENTER,
        spaceAfter=14)

    h2_style = S("VH2",
        fontSize=13, leading=16, textColor=NAVY,
        fontName="Helvetica-Bold",
        spaceBefore=16, spaceAfter=4)

    h3_style = S("VH3",
        fontSize=11, leading=14, textColor=NAVY,
        fontName="Helvetica-Bold",
        spaceBefore=10, spaceAfter=3)

    h4_style = S("VH4",
        fontSize=10.5, leading=13, textColor=DARK,
        fontName="Helvetica-BoldOblique",
        spaceBefore=7, spaceAfter=2)

    body_style = S("VBody",
        fontSize=10, leading=14, textColor=DARK,
        fontName="Helvetica",
        alignment=TA_JUSTIFY, spaceAfter=5)

    bullet_style = S("VBullet",
        fontSize=10, leading=13, textColor=DARK,
        fontName="Helvetica",
        leftIndent=16, bulletIndent=4, spaceAfter=3,
        bulletFontName="Helvetica", bulletFontSize=10)

    fn_style = S("VFn",
        fontSize=8.5, leading=11, textColor=GREY,
        fontName="Helvetica-Oblique",
        spaceAfter=2)

    abstract_style = S("VAbstract",
        fontSize=9.5, leading=13, textColor=DARK,
        fontName="Helvetica-Oblique",
        leftIndent=18, rightIndent=18,
        alignment=TA_JUSTIFY, spaceAfter=8)

    story = []

    def add_hr():
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=RULE, spaceAfter=4))

    # ── identify title ────────────────────────────────────────────────────────
    title_done = False
    abstract_mode = False
    body_blocks = []
    for typ, content in blocks:
        if typ == "h1" and not title_done:
            story.append(Paragraph(clean(content), title_style))
            story.append(Paragraph(
                "VANI System — IEEE Conference Paper",
                sub_style))
            story.append(Paragraph(
                "[Author Name] · [Institution] · [email@domain]",
                author_style))
            story.append(HRFlowable(width="100%", thickness=1,
                                    color=NAVY, spaceAfter=10))
            title_done = True
        else:
            body_blocks.append((typ, content))

    # ── body ─────────────────────────────────────────────────────────────────
    in_abstract = False
    for typ, content in body_blocks:

        if typ == "hr":
            add_hr()
            continue

        if typ == "h2":
            txt = clean(content)
            if "Abstract" in txt:
                in_abstract = True
            else:
                in_abstract = False
            story.append(Paragraph(txt, h2_style))
            continue

        if typ == "h3":
            story.append(Paragraph(clean(content), h3_style))
            continue

        if typ == "h4":
            story.append(Paragraph(clean(content), h4_style))
            continue

        if typ == "para":
            style = abstract_style if in_abstract else body_style
            if "Index Terms" in content or "**Index Terms**" in content:
                in_abstract = False
                story.append(Paragraph(clean(content), fn_style))
            else:
                story.append(Paragraph(clean(content), style))
            continue

        if typ == "bullet":
            story.append(Paragraph(
                u"\u2022\u2002" + clean(content), bullet_style))
            continue

        if typ == "enum":
            story.append(Paragraph(
                u"\u25b8\u2002" + clean(content), bullet_style))
            continue

        if typ == "footnote":
            story.append(Paragraph(clean(content), fn_style))
            continue

        if typ == "table":
            rows = content
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)

            # normalise rows to num_cols
            data = []
            for row in rows:
                padded = list(row) + [""] * (num_cols - len(row))
                data.append([clean(c) for c in padded])

            # Compute column widths
            available = COL_W - 2 * mm
            col_w = available / num_cols
            col_widths = [col_w] * num_cols

            tbl = Table(data, colWidths=col_widths, repeatRows=1)

            # header
            tbl_style = [
                # header row
                ("BACKGROUND",  (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR",   (0, 0), (-1, 0), WHITE),
                ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",    (0, 0), (-1, 0), 8),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 5),
                ("TOPPADDING",  (0, 0), (-1, 0), 5),
                # data rows
                ("FONTNAME",    (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE",    (0, 1), (-1, -1), 8),
                ("TOPPADDING",  (0, 1), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 1), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT]),
                # borders
                ("GRID",        (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ("LINEABOVE",   (0, 0), (-1, 0),  1,    NAVY),
                ("LINEBELOW",   (0, -1), (-1, -1), 0.75, NAVY),
                ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN",       (0, 0), (-1, -1), "LEFT"),
            ]
            tbl.setStyle(TableStyle(tbl_style))
            story.append(KeepTogether([tbl, Spacer(1, 8)]))
            continue

    # ── build ─────────────────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=2.4 * cm,
        bottomMargin=2.0 * cm,
        title="VANI — Indic Speech Intelligence System",
        author="VANI Research",
    )

    def on_page(canvas, doc):
        canvas.saveState()
        # header bar
        canvas.setFillColor(NAVY)
        canvas.rect(MARGIN, PAGE_H - 1.7*cm,
                    PAGE_W - 2*MARGIN, 0.3*cm, fill=1, stroke=0)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(WHITE)
        canvas.drawString(MARGIN + 2, PAGE_H - 1.6*cm,
                          "VANI: Novel Techniques for Indic ASR")
        canvas.drawRightString(PAGE_W - MARGIN, PAGE_H - 1.6*cm,
                               "IEEE Conference Paper — 2026")
        # footer
        canvas.setFillColor(GREY)
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(PAGE_W / 2, 1.2*cm,
                                 f"Page {doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    print(f"[PDF]  Saved: {PDF_OUT}")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"Parsing {SRC}...")
    blocks = parse_md(SRC)
    print(f"  {len(blocks)} blocks parsed")

    print("Building DOCX...")
    build_docx(blocks)

    print("Building PDF...")
    build_pdf(blocks)

    print("\nDone.")
    print(f"  DOCX: {DOCX_OUT}")
    print(f"  PDF:  {PDF_OUT}")
