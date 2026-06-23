#!/usr/bin/env python3
"""
Generate VANI LRP Report as a Word document (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Header ────────────────────────────────────────────────────────────────────
section = doc.sections[0]
header  = section.header
hp      = header.paragraphs[0]
hp.text = "VANI – Voice Analysis & Neural Intelligence  |  Literature Review Proposal  |  SOATE-44"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in hp.runs:
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_spacing(paragraph, space_before=6, space_after=6, line_rule=WD_LINE_SPACING.MULTIPLE, line_val=1.15):
    pf = paragraph.paragraph_format
    pf.space_before    = Pt(space_before)
    pf.space_after     = Pt(space_after)
    pf.line_spacing_rule = line_rule
    pf.line_spacing    = line_val
    pf.alignment       = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title(doc, text, size=16, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold          = bold
    run.italic        = italic
    run.font.size     = Pt(size)
    run.font.name     = "Calibri"
    pf = p.paragraph_format
    pf.space_before   = Pt(space_before)
    pf.space_after    = Pt(space_after)
    return p


def add_heading(doc, text, size=13, bold=True, space_before=14, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(size)
    run.font.name  = "Calibri"
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    return p


def add_subheading(doc, text, size=11.5, bold=True, italic=True, space_before=10, space_after=3):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.name  = "Calibri"
    run.font.color.rgb = RGBColor(0x2E, 0x54, 0x96)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    return p


def add_body(doc, text, size=11, space_before=4, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    set_spacing(p, space_before=space_before, space_after=space_after)
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    return p


def shade_cell(cell, fill_hex="1F3864"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()  # top space

add_title(doc, "VANI – Voice Analysis & Neural Intelligence", size=18, space_before=24, space_after=4)
add_title(doc, "A Literature Review Proposal", size=16, space_before=0, space_after=8)
add_title(doc, "SOATE-44 Project Proposal / Synopsis", size=12, bold=False, italic=True, space_before=0, space_after=40)

# Metadata block
meta = [
    ("Project Code",  "SOATE-44"),
    ("Document Type", "Literature Review Paper (LRP)"),
    ("Date",          "March 2026"),
    ("Classification","RESTRICTED – For Academic Submission Only"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1.  Introduction")

add_body(doc,
    "The Indian subcontinent and its surrounding regions constitute one of the most linguistically diverse "
    "operational theatres in the world. Spanning a contiguous arc that includes South Asia, Central Asia, "
    "and South-East Asia, this zone encompasses scores of languages drawn from four distinct script families: "
    "Devanagari (Hindi, Nepali, Maithili, Dogri), Gurmukhi (Punjabi), Perso-Arabic (Urdu, Kashmiri, Sindhi, "
    "Pashto, Farsi, Arabic), and Latin or logographic systems (Burmese, Tibetan, Chinese, Uzbek, Kazakh, Tajik). "
    "Radio communications in this environment are routinely conducted in any one of these languages, and "
    "often in code-switched speech that mixes two or more. For signals intelligence (SIGINT) analysts operating "
    "in the field, the challenge is not merely transcription but a fully automated pipeline that can identify "
    "an unknown language, transcribe it, translate it into English, and distil the result into a structured "
    "intelligence summary — all in near-real time and without any connectivity to remote servers."
)

add_body(doc,
    "Conventional SIGINT workflows rely on teams of trained linguists supported by cloud-based speech APIs. "
    "These workflows are unsuitable in forward-deployed or communications-denied environments where (a) network "
    "connectivity cannot be guaranteed, (b) sensitive audio cannot be transmitted to external infrastructure, "
    "and (c) compute resources are constrained to commodity hardware with no dedicated graphics processing unit. "
    "The emergence of transformer-based models (Vaswani et al., 2017) has made it theoretically possible to "
    "deploy high-quality automatic speech recognition (ASR), language identification (LangID), machine translation "
    "(MT), and natural language generation (NLG) on a single CPU-bound workstation — but the practical integration "
    "of these components into a coherent, operational pipeline demands careful engineering and a thorough "
    "understanding of each component's limitations."
)

add_body(doc,
    "VANI (Voice Analysis and Neural Intelligence) is a research prototype that addresses precisely this gap. "
    "Designed to operate entirely offline on an 8 GB RAM, CPU-only machine, VANI implements a sequential "
    "eight-stage pipeline: voice activity detection (VAD), audio preprocessing, chunking, ASR via OpenAI "
    "Whisper (Radford et al., 2022), a three-way language identification ensemble combining Whisper's built-in "
    "probability, FastText LID-176 (Joulin et al., 2017), and Facebook MMS-LID-256 (Pratap et al., 2023), "
    "machine translation via NLLB-200 (Costa-jussà et al., 2022) with an IndicTrans2 fallback for Dogri "
    "(Gala et al., 2023), keyword-based threat detection, and automated intelligence summary (ISUM) generation "
    "using either a rule-based module or Qwen2.5. The system targets 19 languages with coverage across all "
    "major script families present in the operational theatre."
)

add_body(doc,
    "This literature review proposal surveys ten foundational and applied works that underpin VANI's design. "
    "The review is organised into five thematic clusters: (1) the Transformer architecture as the common "
    "foundation, (2) ASR approaches for multilingual and low-resource speech, (3) ensemble LangID methods, "
    "(4) neural machine translation with a focus on Indic language coverage, and (5) intelligence extraction "
    "and future enhancement pathways including cross-lingual NER, abstractive summarisation, and speaker "
    "diarisation. For each cluster, the review identifies the research gap addressed, the contribution of "
    "each cited work, and VANI's current or planned integration of that work."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 – PROBLEM STATEMENT & OBJECTIVES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2.  Problem Statement & Objectives")

add_subheading(doc, "2.1  Operational Problem")

add_body(doc,
    "Forward-deployed intelligence personnel require a system that can process raw radio intercept audio "
    "and produce structured intelligence products with minimal human intervention. The operational constraints "
    "that define this problem are severe and non-negotiable:"
)

add_bullet(doc,
    "Language diversity: Intercepts may arrive in any of 19+ languages spanning four script families. "
    "The system must identify the language before it can translate, yet language identification from "
    "short, noisy audio fragments is inherently ambiguous — particularly for closely related language "
    "pairs such as Hindi and Punjabi, or Urdu and Hindi, which share high lexical overlap but are "
    "written in different scripts and are spoken by different communities with distinct operational significance."
)
add_bullet(doc,
    "Network isolation: The system must function with zero internet connectivity. All models must be "
    "stored locally and loaded on demand. Cloud APIs (Google Speech, DeepL, OpenAI API) are categorically "
    "excluded. This eliminates the majority of commercially available speech-to-text and translation tools."
)
add_bullet(doc,
    "Hardware constraint: The target platform is an 8 GB RAM workstation with a contemporary multi-core "
    "CPU but no GPU. Large language models requiring GPU acceleration (e.g., 7B-parameter models) are "
    "excluded unless they can be quantised to fit within the available RAM budget alongside the other "
    "pipeline components loaded simultaneously."
)
add_bullet(doc,
    "Latency: Analysts require outputs within a reasonable time window. A real-time factor (RTF) of less "
    "than 3.0 — meaning processing takes no more than three times the duration of the audio — is the "
    "target ceiling for the full pipeline on a 60-second intercept."
)
add_bullet(doc,
    "Structured output: Raw transcripts and translations are insufficient. Intelligence consumers require "
    "structured 5W summaries (Who, What, Where, When, Assessment), threat level assessments, and "
    "exportable reports in PDF and DOCX formats suitable for inclusion in intelligence products."
)

add_subheading(doc, "2.2  Primary Objectives")

add_body(doc, "The primary objectives of the VANI project are:")

add_bullet(doc,
    "(a)  Multilingual ASR on noisy radio audio: Achieve transcription across 19 target languages using "
    "a single model without per-language fine-tuning, operating within the RAM and CPU constraints."
)
add_bullet(doc,
    "(b)  Robust language identification: Implement a voting ensemble that resolves language ambiguity "
    "with a confidence score, flags uncertain classifications, and correctly identifies Hindi vs. Punjabi "
    "from both script cues and audio characteristics."
)
add_bullet(doc,
    "(c)  Machine translation to English: Route each identified language to the appropriate translation "
    "model, covering all 19 target languages. Produce a back-translation chrF score as an automatic "
    "quality proxy when no reference translation is available."
)
add_bullet(doc,
    "(d)  Automated ISUM generation: Extract the 5W fields, assign a threat level (CRITICAL / HIGH / "
    "MEDIUM / LOW / CLEAR), and flag quality issues (low language confidence, translation failure, "
    "low ASR confidence) in a structured intelligence summary report."
)
add_bullet(doc,
    "(e)  Analyst annotation and training data collection: Provide an interface for analysts to correct "
    "transcripts, translations, and ISUM fields, storing corrections in a structured database for "
    "future model fine-tuning."
)

add_subheading(doc, "2.3  Sub-Objectives and Performance Targets")

add_body(doc,
    "In addition to the primary objectives, the following sub-objectives define measurable success criteria "
    "for the research prototype:"
)

add_bullet(doc, "Real-Time Factor below 3.0 on a 60-second intercept on the target hardware configuration.")
add_bullet(doc, "Threat detection recall above 80% on a held-out keyword set drawn from the configured threat lexicon.")
add_bullet(doc, "5W ISUM completeness score of 3 or above (out of 5) for intercepts with clear speech and known language.")
add_bullet(doc, "Language identification accuracy above 85% on a balanced evaluation set of the 19 target languages.")
add_bullet(doc, "PDF and DOCX report export with metadata, threat badge, transcript, translation, and optional analyst metrics.")
add_bullet(doc, "SQLite database persistence enabling full-text search, threat-level filtering, and date-range queries across all stored intercepts.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 – LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3.  Literature Review")

add_body(doc,
    "The ten papers reviewed in this section are organised into five thematic clusters corresponding to the "
    "major technical components of the VANI pipeline. Each cluster opens with a statement of the research gap "
    "that motivates the selection, followed by a discussion of each paper's contribution, VANI's integration "
    "of that contribution, and the residual gaps that remain open. A consolidated summary table is provided "
    "at the end of this section."
)

# ── Cluster 1 ─────────────────────────────────────────────────────────────────
add_subheading(doc, "3.1  Cluster 1 – Foundational Architecture: The Transformer (P6)")

add_body(doc,
    "Research gap: Prior to 2017, sequence-to-sequence models for speech and text relied predominantly on "
    "recurrent architectures (LSTM, GRU) that processed inputs sequentially. This imposed hard limits on "
    "parallelisation during training and on the modelling of long-range dependencies, both of which are "
    "critical in multilingual speech and translation tasks."
)

add_body(doc,
    "Vaswani et al. (2017) introduced the Transformer architecture in 'Attention Is All You Need', replacing "
    "recurrence entirely with multi-head self-attention mechanisms. The scaled dot-product attention allows "
    "each position in an input sequence to attend to every other position simultaneously, enabling the model "
    "to capture long-range linguistic dependencies that recurrent models struggled to maintain across many "
    "time steps. The encoder-decoder structure with positional encodings became the canonical blueprint for "
    "virtually all subsequent large-scale sequence models. The Transformer also enabled far more efficient "
    "training via data parallelism, which in turn made it feasible to train on the dataset scales required "
    "for multilingual zero-shot transfer."
)

add_body(doc,
    "VANI integration: The Transformer is not merely a cited work but the architectural substrate of every "
    "neural model in the VANI pipeline. Whisper's encoder-decoder processes mel-spectrogram frames; "
    "NLLB-200 uses a transformer encoder-decoder for translation; IndicTrans2 is transformer-based; "
    "MMS-LID uses the wav2vec 2.0 transformer encoder for audio representation; and Qwen2.5 is a "
    "decoder-only causal transformer. Understanding the attention mechanism is therefore prerequisite to "
    "diagnosing any failure mode in VANI. For example, the O(n²) attention complexity means that very long "
    "intercepts — above approximately 30 seconds — must be chunked before ASR to remain within RAM budgets, "
    "a design decision that directly shaped the pipeline's preprocessing stage."
)

add_body(doc,
    "Residual gap: The quadratic complexity of standard attention remains a practical constraint on sequence "
    "length. Linear attention approximations (e.g., Longformer, FlashAttention) are an active research area "
    "but have not yet been incorporated into any of the VANI component models."
)

# ── Cluster 2 ─────────────────────────────────────────────────────────────────
add_subheading(doc, "3.2  Cluster 2 – Automatic Speech Recognition: Whisper and IndicWav2Vec (P1, P7)")

add_body(doc,
    "Research gap: Deploying ASR across 19 diverse languages without per-language fine-tuned models requires "
    "a zero-shot multilingual approach. Existing fine-tuned models achieve excellent WER on their target "
    "language but require labelled in-domain audio for each new language — an impractical requirement for "
    "a system covering 19 languages in a resource-constrained deployment."
)

add_body(doc,
    "Radford et al. (2022) introduced Whisper, an encoder-decoder transformer trained on 680,000 hours of "
    "weakly supervised multilingual and multilingual audio-text pairs sourced from the internet. The model "
    "supports 99 languages zero-shot and was demonstrated to approach or exceed the performance of "
    "supervised baselines on many languages despite never being fine-tuned on labelled in-domain data. "
    "The large-v3-turbo variant used in VANI is a distilled and quantised version that reduces the memory "
    "footprint and inference latency while preserving most of the recognition quality. Whisper's built-in "
    "language detection produces a per-segment probability distribution over 99 languages, which VANI "
    "incorporates as the first of three votes in the LangID ensemble. The initial_prompt mechanism allows "
    "domain vocabulary — military callsigns, weapons designations, grid reference formats — to be injected "
    "into the decoder's context, improving recognition of out-of-vocabulary terms common in tactical communications."
)

add_body(doc,
    "A critical failure mode identified during VANI development is Whisper's tendency to misidentify Punjabi "
    "(Gurmukhi script) as Hindi (Devanagari script). Because Punjabi and Hindi share substantial phonemic "
    "overlap, Whisper's audio-based language detection alone is insufficient to discriminate reliably between "
    "them. This motivated the introduction of script-aware post-processing and the three-way LangID ensemble "
    "described in Section 3.3."
)

add_body(doc,
    "Javed et al. (2022) presented IndicWav2Vec, a wav2vec 2.0 model continued-pretrained on 17,000 hours of "
    "Indian-language speech covering nine Indic languages. IndicWav2Vec achieves state-of-the-art word error "
    "rates on the IndicSUPERB benchmark, with reported WER of approximately 22.3% for Punjabi under standard "
    "evaluation conditions. This figure serves as an empirical performance floor for VANI's ASR module: if "
    "VANI's zero-shot Whisper exceeds 22.3% WER on Punjabi, it is operating below the achievable benchmark "
    "for that language. The paper also provides WER baselines for Hindi, Bengali, Nepali, and other VANI "
    "target languages, enabling the METRICS tab to contextualise analyst-entered reference WER values."
)

add_body(doc,
    "VANI integration: Whisper (large-v3-turbo, CTranslate2 int8 quantisation) serves as the sole ASR model. "
    "The configuration parameters beam_size=4, temperature=0.0, condition_on_previous_text=False, and "
    "no_speech_threshold=0.70 were set to balance accuracy, determinism, and noise robustness. IndicWav2Vec "
    "currently provides baseline targets only; it is identified as a candidate for domain-adaptive fine-tuning "
    "once 500 or more annotated intercepts per language have been accumulated via the annotation subsystem."
)

add_body(doc,
    "Residual gap: Whisper's zero-shot WER on heavily noise-degraded radio audio — with channel compression, "
    "squelch artefacts, and overlapping transmissions — has not been systematically benchmarked. The "
    "no_speech_threshold parameter mitigates some noise segments but cannot recover intelligibility from "
    "heavily corrupted audio. Domain-specific fine-tuning of Whisper on labelled tactical communications "
    "audio remains a priority for future work."
)

# ── Cluster 3 ─────────────────────────────────────────────────────────────────
add_subheading(doc, "3.3  Cluster 3 – Language Identification: MMS-LID and FastText (P4, P5)")

add_body(doc,
    "Research gap: Single-modality language identification — whether audio-based or text-based — is "
    "insufficient when dealing with short, noisy speech fragments from closely related language pairs. "
    "A robust system must triangulate across multiple independent evidence sources, weight them by "
    "confidence, and detect cases where the evidence is too ambiguous to assign a label reliably."
)

add_body(doc,
    "Pratap et al. (2023) presented the Massively Multilingual Speech (MMS) project, scaling speech "
    "technology to over 1,000 languages using the New Testament audio corpus as a low-resource training "
    "resource. The MMS-LID-256 model, based on wav2vec 2.0 fine-tuned for language identification, achieves "
    "over 90% accuracy on 256 languages from raw audio alone, without requiring any prior transcription. "
    "Its audio-based representation means it captures prosodic and phonotactic features that text-based "
    "models cannot access. For the Hindi/Punjabi discrimination problem, MMS-LID can exploit differences "
    "in retroflex consonant distribution, vowel length, and tonal patterns that are phonemically distinctive "
    "even if the lexicon overlaps substantially."
)

add_body(doc,
    "Joulin et al. (2017) described FastText's efficient text classification approach using character-level "
    "n-gram features combined with a shallow neural architecture. The lid.176.bin model identifies 176 "
    "languages from text input with reported accuracy above 97% and sub-millisecond inference latency. "
    "FastText's character n-gram features are particularly powerful for script discrimination: Gurmukhi "
    "(Punjabi), Devanagari (Hindi/Nepali), and Perso-Arabic (Urdu/Kashmiri/Sindhi) produce entirely "
    "disjoint n-gram vocabularies, making script-level separation nearly perfect. The limitation is that "
    "very short text fragments — fewer than 20 tokens — produce lower-confidence predictions, which is "
    "common when early ASR chunks are short."
)

add_body(doc,
    "VANI integration: The three-way ensemble combines (1) Whisper's language probability from the initial "
    "audio chunk, (2) FastText's LangID score on the ASR transcript, and (3) MMS-LID's audio-based "
    "classification. A confidence-weighted voting scheme is applied: unanimous agreement yields the "
    "majority-confidence label; two-way agreement yields the average confidence of the agreeing pair; "
    "three-way disagreement falls back to the single highest-confidence vote. Any result below the "
    "CONFIDENCE_THRESHOLD of 0.60 is flagged as language_uncertain in the ISUM output, triggering "
    "the LOW_LANG_CONFIDENCE quality flag. A specific Punjabi override rule — triggered when FastText "
    "or MMS assigns pa with confidence ≥0.55 while Whisper assigns hi — forces the language to Punjabi "
    "and routes translation via NLLB-200 rather than IndicTrans2, resolving a previously identified "
    "DynamicCache crash."
)

add_body(doc,
    "Residual gap: MMS-LID's performance degrades with heavy radio channel noise, and FastText's confidence "
    "drops on short fragments. Neither model has been evaluated specifically on tactical radio communications "
    "audio. Collecting labelled radio intercepts for ensemble calibration is an identified priority."
)

# ── Cluster 4 ─────────────────────────────────────────────────────────────────
add_subheading(doc, "3.4  Cluster 4 – Machine Translation: NLLB-200 and IndicTrans2 (P2, P3)")

add_body(doc,
    "Research gap: No single publicly available translation model covers all 19 VANI target languages with "
    "acceptable quality while fitting within an 8 GB RAM budget. In particular, Dogri (doi) — a scheduled "
    "Indian language spoken predominantly in Jammu — is absent from NLLB-200's language set despite being "
    "operationally significant."
)

add_body(doc,
    "Costa-jussà et al. (2022) presented NLLB-200 (No Language Left Behind), a 200-language neural machine "
    "translation model developed by Meta AI. The distilled 600M parameter variant achieves a +44% improvement "
    "in BLEU over the prior best publicly available multilingual MT model on the FLORES-200 benchmark. "
    "For the 18 of 19 VANI target languages present in NLLB-200's language set, the distilled-600M model "
    "fits comfortably within the RAM budget and delivers translation quality sufficient for intelligence "
    "triage purposes — confirming the core message of intercepts even if stylistic accuracy is imperfect. "
    "NLLB-200 also supports back-translation (translating from English back to the source language) enabling "
    "VANI to compute a reference-free translation quality estimate as the chrF score between the "
    "back-translation and the original transcript."
)

add_body(doc,
    "Gala et al. (2023) presented IndicTrans2, a translation model specifically designed to cover all 22 "
    "scheduled Indian languages, including Dogri (doi), which is the critical gap in NLLB-200's coverage. "
    "IndicTrans2 outperforms NLLB-200 on 18 of 22 Indic language pairs on the IN22 benchmark. However, "
    "the model was trained against the transformers library at version 4.x, and multiple breaking changes "
    "in transformers 5.x required significant compatibility engineering before VANI could use the model. "
    "Specifically, the removal of the transformers.onnx submodule required a stub import in the model's "
    "configuration file; the tie_weights() method required **kwargs support; the lang_code_to_id attribute "
    "was replaced by convert_tokens_to_ids(); the _switch_to_input_mode() method must be explicitly called "
    "before tokenisation; and use_cache=False is required in generate() to avoid DynamicCache / "
    "EncoderDecoderCache incompatibilities. All of these fixes have been applied and are documented in "
    "VANI's translation module."
)

add_body(doc,
    "VANI integration: Language routing logic directs Dogri (doi) exclusively to the IndicTrans2 route, "
    "all other non-English languages to the NLLB-200 route, and English directly to the output without "
    "translation. The routing is deterministic and based on the final_language value established by the "
    "LangID ensemble. Back-translation is enabled for NLLB-200 languages as a Tier 1 automatic quality metric."
)

add_body(doc,
    "Residual gap: Kashmiri (ks) and Sindhi (sd) NLLB-200 quality scores are lower than for Hindi/Punjabi/Urdu "
    "due to limited training data for those languages. Tibetan (bo) and Burmese (my) are covered by NLLB-200 "
    "but at lower quality than the Indic languages. A systematic evaluation of VANI's translation quality "
    "across all 19 languages using the FLORES-200 benchmark prompts remains outstanding."
)

# ── Cluster 5 ─────────────────────────────────────────────────────────────────
add_subheading(doc, "3.5  Cluster 5 – Intelligence Extraction & Future Enhancement (P8, P9, P10)")

add_body(doc,
    "Research gap: VANI's current ISUM module uses regular expressions and keyword matching for entity "
    "extraction and a rule-based 5W template for summary generation. This approach is brittle: it fails "
    "on paraphrased expressions, cannot handle multi-party conversations, and lacks the abstractive "
    "capabilities needed to synthesise information across multiple transcript segments. Three complementary "
    "research directions address these gaps."
)

add_body(doc,
    "Conneau et al. (2020) presented XLM-RoBERTa (XLM-R), a cross-lingual language model trained via masked "
    "language modelling on 2.5 terabytes of filtered CommonCrawl text in 100 languages. XLM-R demonstrates "
    "strong zero-shot cross-lingual transfer: a model fine-tuned on English named entity recognition (NER) "
    "annotations achieves competitive performance on the same task in other languages without any target-language "
    "labels. This capability is directly relevant to VANI's keyword and entity detection component: rather "
    "than maintaining hand-curated keyword lists in each of 19 languages, a single XLM-R model fine-tuned "
    "on English military NER annotations (weapons, units, locations, persons, events) could be applied "
    "zero-shot to translated or original-language text. This would improve recall on paraphrased threat "
    "indicators and provide structured entity spans for the 5W ISUM fields."
)

add_body(doc,
    "Zhang et al. (2020) introduced PEGASUS, a sequence-to-sequence model pre-trained with a Gap Sentence "
    "Generation (GSG) objective that masks entire key sentences from a document and trains the model to "
    "regenerate them — a pre-training task that is semantically equivalent to summarisation. PEGASUS achieves "
    "state-of-the-art performance on 12 summarisation benchmarks and demonstrates effective few-shot "
    "fine-tuning with as few as 1,000 examples. The relevance to VANI is direct: the ISUM task — compressing "
    "a raw intercept transcript into a structured 5W intelligence summary — is an instance of domain-specific "
    "abstractive summarisation. The annotation system built into VANI's [A] tab is specifically designed to "
    "accumulate transcript–ISUM pairs that could be used to fine-tune a multilingual PEGASUS variant. The "
    "current Qwen2.5 LLM mode is an interim solution; mPEGASUS or a fine-tuned seq2seq model is identified "
    "as a longer-term replacement once sufficient training data is available."
)

add_body(doc,
    "Bredin et al. (2020) presented pyannote.audio, a neural speaker diarisation toolkit providing building "
    "blocks for voice activity detection, speaker change detection, and speaker embedding. The system achieves "
    "a diarisation error rate (DER) below 10% on the AMI and DIHARD benchmarks. Speaker diarisation is a "
    "critical capability gap in VANI's current architecture: all segments from a multi-speaker intercept are "
    "merged into a single undifferentiated transcript, losing the conversational structure. In a tactical "
    "radio intercept, knowing which transmission was made by which station (e.g., callsign Alpha vs. Bravo) "
    "is intelligence-relevant. Integrating pyannote.audio in a planned Phase 5 enhancement would enable "
    "per-speaker ISUM fields, voiceprint tracking across multiple intercepts from the same source, and "
    "a richer conversational structure for the LLM ISUM generator."
)

add_body(doc,
    "VANI integration: XLM-R and PEGASUS are currently referenced as theoretical foundations for the "
    "annotation system's data collection strategy and the future Phase 5 roadmap. pyannote.audio is "
    "explicitly scheduled for Phase 5. The annotation subsystem stores corrections in a schema designed "
    "to produce HuggingFace-compatible training datasets exportable in JSON format compatible with "
    "HuggingFace Trainer, Axolotl, and TRL fine-tuning frameworks."
)

add_body(doc,
    "Residual gap: pyannote.audio performance degrades under high speaker overlap and radio channel noise "
    "conditions, which are precisely the conditions VANI must handle. XLM-R's zero-shot transfer quality "
    "for low-resource Indic languages (Dogri, Kashmiri, Maithili) has not been benchmarked on military NER "
    "tasks. mPEGASUS does not natively support Indic script output, requiring a transliteration or "
    "cross-lingual fine-tuning step before it can generate non-English summaries."
)

# ── Summary Table ──────────────────────────────────────────────────────────────
add_subheading(doc, "3.6  Summary Table of Reviewed Literature", italic=False)

add_body(doc, "Table 1 provides a consolidated overview of the ten reviewed papers.", space_after=4)

table_data = [
    ["P1", "Radford et al.", "2022", "Robust Speech Recognition via Large-Scale Weak Supervision", "ASR", "Whisper ASR backbone; language probability = vote 1 in LangID ensemble"],
    ["P2", "Costa-jussà et al.", "2022", "No Language Left Behind: Scaling Human-Centered MT", "MT", "Primary translation engine for 18/19 target languages; back-translation chrF"],
    ["P3", "Gala et al.", "2023", "IndicTrans2: High-Quality MT for 22 Scheduled Indian Languages", "MT", "Dogri-only fallback translation route; transformers 5.x patches applied"],
    ["P4", "Pratap et al.", "2023", "Scaling Speech Technology to 1,000+ Languages (MMS)", "LangID", "Audio-based vote 3 in LangID ensemble; Hindi/Punjabi disambiguation from audio"],
    ["P5", "Joulin et al.", "2017", "Bag of Tricks for Efficient Text Classification (FastText)", "LangID", "Text-based vote 2 in LangID ensemble; Gurmukhi/Arabic n-gram script detection"],
    ["P6", "Vaswani et al.", "2017", "Attention Is All You Need (Transformer)", "Architecture", "Architectural foundation for all neural models in VANI pipeline"],
    ["P7", "Javed et al.", "2022", "IndicWav2Vec: Multilingual Speech Model for Indian Languages", "ASR", "Provides WER baselines (e.g., Punjabi 22.3%) as performance targets"],
    ["P8", "Conneau et al.", "2020", "Unsupervised Cross-Lingual Representation Learning at Scale (XLM-R)", "NER/NLU", "Future Phase 5: replace regex keyword detection with zero-shot cross-lingual NER"],
    ["P9", "Zhang et al.", "2020", "PEGASUS: Pre-training with Extracted Gap-Sentences for Abstractive Summarisation", "NLG", "Academic foundation for ISUM generation; annotation system designed for fine-tuning"],
    ["P10", "Bredin et al.", "2020", "pyannote.audio: Neural Building Blocks for Speaker Diarisation", "Diarisation", "Planned Phase 5: per-speaker ISUM, voiceprint tracking across intercepts"],
]

headers = ["Paper ID", "Authors", "Year", "Title (Shortened)", "Domain", "VANI Relevance"]

tbl = doc.add_table(rows=1 + len(table_data), cols=6)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Column widths (total ~15.5 cm)
col_widths = [Cm(1.1), Cm(2.4), Cm(1.0), Cm(4.4), Cm(2.0), Cm(5.2)]

# Header row
hdr_cells = tbl.rows[0].cells
for i, (cell, header_text) in enumerate(zip(hdr_cells, headers)):
    shade_cell(cell, "1F3864")
    cell.width = col_widths[i]
    p = cell.paragraphs[0]
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Data rows
for row_idx, row_data in enumerate(table_data):
    row_cells = tbl.rows[row_idx + 1].cells
    fill = "DCE6F1" if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, (cell, cell_text) in enumerate(zip(row_cells, row_data)):
        shade_cell(cell, fill)
        cell.width = col_widths[col_idx]
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"
        if col_idx == 0:
            run.bold = True
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after  = Pt(2)

add_body(doc, "Table 1. Summary of the ten reviewed papers organised by contribution domain and VANI relevance.", space_before=4)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 – EXPECTED FINDINGS & RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4.  Expected Findings & Recommendations")

add_subheading(doc, "4.1  Expected Performance Findings")

add_body(doc,
    "Based on the empirical benchmarks reported in the reviewed literature and on preliminary results "
    "observed during VANI's development, the following performance findings are expected when the system "
    "is evaluated against held-out intercept audio:"
)

add_body(doc,
    "ASR accuracy: Whisper large-v3-turbo is expected to achieve WER below 30% for Hindi (hi) and Punjabi "
    "(pa) — the two highest-resource Indic languages in its training data — under clear-channel conditions. "
    "This target is calibrated against the IndicWav2Vec baselines from Javed et al. (2022), specifically "
    "the reported 22.3% WER for Punjabi under standard evaluation. VANI's zero-shot Whisper is not expected "
    "to match fine-tuned IndicWav2Vec on Punjabi in isolation, but its ability to handle all 19 languages "
    "without per-language models is the operationally relevant advantage. For lower-resource languages "
    "(Dogri, Kashmiri, Maithili, Tibetan), WER is expected to be substantially higher, in the 40–60% range, "
    "reflecting Whisper's sparse training coverage for those languages."
)

add_body(doc,
    "Translation quality: NLLB-200 distilled-600M is expected to achieve BLEU scores in the 25–45 range "
    "for Hindi-English and Punjabi-English translation on FLORES-200 prompts, consistent with the +44% "
    "BLEU improvement over prior baselines reported by Costa-jussà et al. (2022). For Pashto (ps) and "
    "Tibetan (bo) — languages with smaller NLLB-200 training sets — BLEU scores in the 15–25 range are "
    "expected. IndicTrans2 is expected to produce higher-quality Dogri translation than any NLLB-200 "
    "equivalent, but the absence of a publicly available Dogri reference translation corpus makes formal "
    "BLEU evaluation impractical at this stage. Back-translation chrF scores above 0.45 on Hindi and "
    "Punjabi are considered indicative of acceptable translation quality for intelligence triage purposes."
)

add_body(doc,
    "ISUM completeness: The rule-based ISUM module is expected to achieve a 5W completeness score of "
    "3 out of 5 or higher for intercepts containing at least one explicit location reference, one time "
    "reference, and callsign-style identifiers (Alpha/Bravo patterns). For intercepts that lack these "
    "structural cues, the rule-based ISUM will populate fewer fields, and the completeness score will "
    "reflect this. The Qwen2.5 LLM ISUM mode is expected to improve completeness for paraphrased or "
    "contextually implicit 5W information, but its quality is bounded by the absence of domain-specific "
    "fine-tuning data."
)

add_subheading(doc, "4.2  Operational Findings")

add_body(doc,
    "Real-time factor: Preliminary timing measurements indicate that the full VANI pipeline — VAD, "
    "preprocessing, chunking, Whisper ASR (CTranslate2 int8), FastText LangID, MMS-LID, NLLB-200 "
    "translation, keyword detection, and rule-based ISUM — completes a 60-second intercept in under "
    "180 seconds on the target 8 GB RAM CPU-only machine, achieving the RTF <3.0 target. Translation "
    "is the most computationally intensive non-ASR stage, consuming approximately 30–40% of total "
    "processing time for long intercepts. CTranslate2 int8 quantisation of Whisper is the single most "
    "important optimisation enabling the RTF target to be met."
)

add_body(doc,
    "LangID ensemble effectiveness: The three-way LangID ensemble is expected to resolve over 90% of "
    "Hindi/Punjabi ambiguity cases correctly when both FastText's Gurmukhi script detection and MMS-LID's "
    "audio-based Punjabi classification agree against Whisper's Hindi identification. The ensemble "
    "architecture — motivated by the complementary failure modes of Whisper (phonemic similarity) and "
    "FastText (short fragment confidence) described in Pratap et al. (2023) and Joulin et al. (2017) — "
    "is the primary mitigation for the most operationally significant misidentification risk in the "
    "target language set."
)

add_body(doc,
    "Database and search: The SQLite database is expected to support full-text search queries over "
    "stored transcripts and translations with sub-second latency for corpora up to 10,000 intercepts, "
    "well within the expected operational scale of a single forward-deployed system."
)

add_subheading(doc, "4.3  Identified Gaps")

add_body(doc,
    "Three structural gaps in VANI's current architecture are directly surfaced by the literature review:"
)

add_bullet(doc,
    "Speaker diarisation is absent. All transmissions in a multi-party intercept are merged into a single "
    "transcript, losing the conversational structure. This means VANI cannot distinguish between transmission "
    "sources, cannot attribute statements to individual callsigns, and cannot track whether a specific "
    "voice has appeared in previous intercepts. Bredin et al. (2020) demonstrate that neural diarisation "
    "with DER <10% is achievable, but performance under radio channel noise conditions has not been "
    "established."
)
add_bullet(doc,
    "Keyword detection via regular expressions is brittle and recall-limited. The current implementation "
    "matches surface-form keywords and does not capture paraphrases, synonyms, or contextual threat "
    "indicators. Conneau et al. (2020) demonstrate that XLM-R zero-shot NER transfer from English to "
    "other languages can substantially improve structured entity recall, but this requires fine-tuning "
    "data in the military NER domain which is not yet available."
)
add_bullet(doc,
    "ISUM generation quality is bounded by the absence of fine-tuning data. The rule-based ISUM is "
    "reliable but inflexible; the Qwen2.5 LLM ISUM is flexible but produces inconsistent output quality "
    "without domain-specific instruction tuning. Zhang et al. (2020) establish that as few as 1,000 "
    "annotated examples are sufficient for high-quality fine-tuning of a summarisation model, making "
    "the annotation subsystem's data collection function a critical path item."
)

add_subheading(doc, "4.4  Recommendations")

add_body(doc, "The following recommendations are made for VANI's development roadmap:")

add_bullet(doc,
    "(a)  Phase 5 speaker diarisation: Integrate pyannote.audio (Bredin et al., 2020) to enable "
    "per-speaker segment labelling, per-speaker 5W ISUM fields, and voiceprint tracking. Evaluation "
    "should include a noise-augmented test set with simulated radio channel conditions. The planned "
    "integration should use pyannote's offline-capable speaker embedding components to preserve the "
    "system's network-independence constraint."
)
add_bullet(doc,
    "(b)  IndicWav2Vec domain fine-tuning: Once 500 or more annotated intercepts per language have been "
    "accumulated via the annotation subsystem, fine-tune IndicWav2Vec (Javed et al., 2022) on the "
    "collected ASR corrections. Evaluate WER improvement against the Whisper zero-shot baseline for "
    "each Indic language. This is expected to produce the largest accuracy gains for Punjabi, Hindi, "
    "and Bengali given IndicWav2Vec's pre-training coverage of those languages."
)
add_bullet(doc,
    "(c)  XLM-R military NER: Define a military NER annotation schema covering entities relevant to "
    "the ISUM 5W fields (PERSON/callsign, LOCATION/grid-ref, WEAPON, UNIT, EVENT, TIME). Annotate "
    "at minimum 500 English-language training examples drawn from open-source military communication "
    "transcripts. Fine-tune XLM-R-base (Conneau et al., 2020) on these annotations and evaluate "
    "zero-shot transfer quality on Punjabi and Hindi VANI outputs. Replace the regex keyword detection "
    "module with the XLM-R NER module if F1 exceeds 0.70 on the held-out evaluation set."
)
add_bullet(doc,
    "(d)  mPEGASUS ISUM fine-tuning: Once 1,000 or more transcript-to-ISUM annotation pairs have been "
    "collected across the target languages, evaluate fine-tuning a multilingual PEGASUS or BART-based "
    "model (Zhang et al., 2020) as an alternative to the Qwen2.5 LLM mode. A seq2seq model fine-tuned "
    "on domain-specific summarisation data is expected to produce more consistent and verifiable ISUM "
    "output than a general-purpose instruction-tuned LLM at this scale."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 – REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5.  References")

references = [
    ("Bredin, H., Yin, R., Coria, J. M., Gelly, G., Korshunov, P., Lavechin, M., Fustes, D., Lancelot, H., "
     "Mansoor, W., & Brunet, M.-P. (2020). pyannote.audio: Neural building blocks for speaker diarization. "
     "In Proceedings of the IEEE International Conference on Acoustics, Speech and Signal Processing "
     "(ICASSP). https://doi.org/10.1109/ICASSP40776.2020.9052974"),
    ("Conneau, A., Khandelwal, K., Goyal, N., Chaudhary, V., Wenzek, G., Guzmán, F., Grave, E., Ott, M., "
     "Zettlemoyer, L., & Stoyanov, V. (2020). Unsupervised cross-lingual representation learning at scale. "
     "In Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics (ACL) "
     "(pp. 8440–8451). https://doi.org/10.18653/v1/2020.acl-main.747"),
    ("Costa-jussà, M. R., Cross, J., Çelebi, O., Guzman, F., Heffernan, K., Hou, C., Mahiuddin Chowdhury, "
     "S., Mourachko, A., Sadagopan, N., Schwenk, H., Tan, J., & Tran, C. (2022). No language left behind: "
     "Scaling human-centered machine translation. arXiv preprint arXiv:2207.04672."),
    ("Gala, J., Chitale, P. A., Raghavan, A. K., Gumma, V., Doddapaneni, S., Sai, A., Kunchukuttan, A., "
     "Puduppully, R., Dabre, R., Murthy, R., & Khapra, M. M. (2023). IndicTrans2: Towards high-quality "
     "and accessible machine translation for all 22 scheduled Indian languages. arXiv preprint arXiv:2305.16307."),
    ("Javed, T., Doddapaneni, S., Raman, A., Bhogale, K. S., Ramesh, G., Kunchukuttan, A., Kumar, P., & "
     "Khapra, M. M. (2022). IndicWav2Vec: A multilingual speech model for Indian languages. In "
     "Proceedings of Interspeech 2022. arXiv preprint arXiv:2111.03945."),
    ("Joulin, A., Grave, E., Bojanowski, P., & Mikolov, T. (2017). Bag of tricks for efficient text "
     "classification. In Proceedings of the 15th Conference of the European Chapter of the Association "
     "for Computational Linguistics (EACL) (Vol. 2, pp. 427–431). arXiv preprint arXiv:1607.01759."),
    ("Pratap, V., Tjandra, A., Shi, B., Tomasello, P., Babu, A., Kulkarni, S., Chung, H.-S., Chen, M.-A., "
     "Conneau, A., Subramanian, K., Singh, K., Jain, N., Xu, Y., Ng, C., Singhal, S., Mallya, A., Tu, J., "
     "Saraf, R., Galvez, R., … Auli, M. (2023). Scaling speech technology to 1,000+ languages. arXiv "
     "preprint arXiv:2305.13516."),
    ("Radford, A., Kim, J. W., Xu, T., Brockman, G., McLeavey, C., & Sutskever, I. (2022). Robust speech "
     "recognition via large-scale weak supervision. arXiv preprint arXiv:2212.04356."),
    ("Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & "
     "Polosukhin, I. (2017). Attention is all you need. In Advances in Neural Information Processing "
     "Systems (NeurIPS) (Vol. 30). arXiv preprint arXiv:1706.03762."),
    ("Zhang, J., Zhao, Y., Saleh, M., & Liu, P. J. (2020). PEGASUS: Pre-training with extracted "
     "gap-sentences for abstractive summarization. In Proceedings of the 37th International Conference "
     "on Machine Learning (ICML) (pp. 11328–11339). arXiv preprint arXiv:1912.08777."),
]

for ref_text in references:
    p = doc.add_paragraph()
    run = p.add_run(ref_text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    pf = p.paragraph_format
    pf.space_before       = Pt(2)
    pf.space_after        = Pt(5)
    pf.left_indent        = Cm(1.0)
    pf.first_line_indent  = Cm(-1.0)
    pf.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 – DECLARATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6.  Declaration")

add_body(doc,
    "I hereby declare that this Literature Review Proposal has been prepared independently for the purposes "
    "of the SOATE-44 course project submission. The work presented is my own original research and analysis "
    "except where explicitly attributed to the cited sources. All referenced works have been properly "
    "acknowledged in accordance with APA 7th edition citation standards. This document has not been submitted, "
    "in whole or in part, to any other course, institution, or publication."
)

p_name = doc.add_paragraph()
p_name.paragraph_format.space_before = Pt(18)
p_name.paragraph_format.space_after  = Pt(4)
run_label = p_name.add_run("Name:  ")
run_label.bold = True
run_label.font.size = Pt(11)
run_label.font.name = "Calibri"
run_blank = p_name.add_run("_" * 40 + "  [Participant Name]")
run_blank.font.size = Pt(11)
run_blank.font.name = "Calibri"

p_inst = doc.add_paragraph()
p_inst.paragraph_format.space_after = Pt(4)
run_label2 = p_inst.add_run("Institution:  ")
run_label2.bold = True
run_label2.font.size = Pt(11)
run_label2.font.name = "Calibri"
run_blank2 = p_inst.add_run("_" * 40)
run_blank2.font.size = Pt(11)
run_blank2.font.name = "Calibri"

p_date = doc.add_paragraph()
p_date.paragraph_format.space_after = Pt(4)
run_label3 = p_date.add_run("Date:  ")
run_label3.bold = True
run_label3.font.size = Pt(11)
run_label3.font.name = "Calibri"
run_blank3 = p_date.add_run("_" * 30)
run_blank3.font.size = Pt(11)
run_blank3.font.name = "Calibri"

p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(20)
p_sig.paragraph_format.space_after  = Pt(4)
run_label4 = p_sig.add_run("Signature:  ")
run_label4.bold = True
run_label4.font.size = Pt(11)
run_label4.font.name = "Calibri"
run_blank4 = p_sig.add_run("_" * 40)
run_blank4.font.size = Pt(11)
run_blank4.font.name = "Calibri"

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = "/Users/vik/offline_ai_system_v2/VANI_LRP_Report.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
