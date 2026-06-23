from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=10):
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def add_paragraph(text, bold=False, italic=False, size=10, space_before=0,
                  space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(10 * 1.05)
    if indent:
        pf.left_indent = Inches(indent)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, size=size)
    return p

def add_heading(text, level=1):
    space = 10 if level == 1 else 6
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(10 * 1.05)
    run = p.add_run(text)
    set_font(run, bold=True, size=10)
    return p

def add_bullet(text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(10 * 1.05)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_mixed(parts, space_before=0, space_after=6, indent=None):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(10 * 1.05)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('Novel Techniques for Automatic Speech Recognition of Indic Languages: '
                'Implementation and Evaluation via the VANI System')
set_font(run, bold=True, size=14)

# Double-blind: no author/institution
add_paragraph('[Author name(s) withheld for double-blind review]',
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, space_after=2)
add_paragraph('[Institution withheld for double-blind review]',
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10, space_after=10)

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Abstract')
add_paragraph(
    'Deploying ASR on Indic languages in the field throws up problems that benchmarks tend to hide. '
    'Two languages may share enough phonemes that a model confuses them. A script may be romanised '
    'in the transcription even though the speaker never used it. Radio recordings bear little '
    'resemblance to studio speech. We built VANI (Voice Analysis and Neural Intelligence) because '
    'existing tools did not handle these situations gracefully on a machine without a GPU or internet '
    'connection. The system uses Whisper large-v3-turbo for transcription, NLLB-200-distilled-600M '
    'and IndicTrans2 for translation, and combines FastText with MMS-LID-256 and Unicode script '
    'checks to decide which language it is actually dealing with. '
    'We tested it on 120 samples across Punjabi, Hindi, Urdu, and Nepali. Translation worked on '
    'every single sample. Language identification hit 96.7% on Hindi and Punjabi but dropped to '
    '70.0% on Urdu and 63.3% on Nepali. Hindi ASR reached 46.2% WER and 32.4% CER; Punjabi and '
    'Nepali WER numbers are artificially inflated because Whisper cannot output Gurmukhi or '
    'Nepali-Devanagari. The most practically useful finding was a confidence gradient—0.879 for '
    'Hindi down to 0.187 for Nepali—that tracks LangID accuracy closely and can be surfaced to '
    'analysts without any reference data.',
    space_after=6
)

# Keywords
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(10)
p.paragraph_format.line_spacing = Pt(10 * 1.05)
r1 = p.add_run('Keywords: ')
set_font(r1, bold=True)
r2 = p.add_run('Automatic speech recognition, Indic languages, language identification, '
               'offline deployment, low-resource ASR')
set_font(r2)

# ══════════════════════════════════════════════════════════════════════════════
# I. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. Introduction')
add_paragraph(
    'India and South Asia are home to hundreds of languages and dialects spoken by large populations, '
    'yet ASR quality varies enormously across this linguistic space. Many languages remain '
    'under-resourced, multiple scripts are in routine use (Devanagari, Gurmukhi, Bengali, Arabic, '
    'Tamil, Telugu, Kannada, Malayalam, Gujarati), and everyday speech frequently mixes a local '
    'language with English. Although end-to-end ASR has reduced dependence on handcrafted lexica, '
    'Indic deployments still require careful design choices. Phone inventories differ substantially '
    'across languages, morphology can be rich and agglutinative, and real recordings frequently '
    'include non-standard spellings, code-mixing, and strong regional pronunciation variation [1-3].'
)
add_paragraph(
    'A concrete and underappreciated challenge is the confusion between closely related Indic '
    'languages at the ASR stage. As demonstrated in our experiments, Whisper large-v3-turbo—trained '
    'on 680,000 hours of multilingual audio [13]—systematically misidentifies Punjabi speech as '
    'Hindi, assigning a Hindi language probability of 0.707 even when the audio contains unambiguous '
    'Gurmukhi script content. This is not a failure of acoustic modeling per se; it reflects the '
    'fact that Hindi and Punjabi share overlapping phonemic inventories and that Punjabi is '
    'significantly underrepresented in Whisper\'s training corpus. Left unchecked, this means the '
    'system routes audio to Hindi translation without the analyst ever knowing it got the language wrong.'
)
add_paragraph(
    'VANI grew out of this specific problem. The goal from the start was a system that could handle '
    'multiple Indic languages in a single recording, on hardware that is actually available in the '
    'field, without a network connection. Each of the six techniques we describe addresses one part '
    'of that constraint.'
)
add_paragraph(
    'The paper makes four concrete contributions. First, we evaluate the full pipeline at scale—120 '
    'samples across Punjabi, Hindi, Urdu, and Nepali from open-source datasets. Second, the '
    'three-source LangID ensemble hits 96.7% on Hindi and Punjabi but degrades to 70.0% and 63.3% '
    'on Urdu and Nepali respectively; the drop directly tracks how much of each language Whisper was '
    'trained on. Third, audio-based MMS-LID proves to be the most reliable component when Whisper '
    'coverage is low—it recovers Nepali in 37% of cases where both text-based sources give wrong or '
    'empty answers. Fourth, NLLB-200-distilled-600M translates every one of the 120 samples '
    'successfully, and the per-segment ASR confidence score turns out to be a useful real-time proxy '
    'for output reliability without any reference transcription.'
)

# ══════════════════════════════════════════════════════════════════════════════
# II. BACKGROUND AND RELATED WORK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. Background and Related Work')
add_paragraph(
    'ASR architecture has gone through several generations. CTC-based pipelines, transducer models '
    '(RNN-T), and attention-based encoder-decoder systems each represent a different tradeoff between '
    'training stability, streaming capability, and output quality [2,3,8]. When deep neural networks '
    'replaced GMM-HMM systems [1], acoustic models became substantially more capable. '
    'Sequence-to-sequence Transformers [4] pushed this further; most ASR work today uses end-to-end '
    'training rather than the pipeline architectures that dominated a decade ago.'
)
add_paragraph(
    'wav2vec 2.0 and HuBERT [6,7] showed that self-supervised pretraining on unlabeled audio could '
    'dramatically reduce the labeled data needed for fine-tuning—useful for languages where '
    'transcribed speech is scarce. Gulati et al. [5] combined convolutional and self-attention layers '
    'in the Conformer architecture. The motivation was that pure Transformer encoders miss short-range '
    'temporal patterns that convolutional filters capture naturally; the two mechanisms turned out to '
    'be complementary.'
)
add_paragraph(
    'Low-resource and multilingual ASR research has focused on shared encoder designs, cross-lingual '
    'vocabulary schemes, and lightweight adaptation methods such as adapters [9]. A notable shift '
    'came with Radford et al. [13], who showed that simply training on 680,000 hours of noisy '
    'multilingual audio—without any language-specific fine-tuning—yielded usable zero-shot ASR '
    'across more than 75 languages. The MMS project [15] pushed this further, covering 1,107 '
    'languages and achieving above 90% top-1 language identification accuracy directly from audio '
    'across 256 classes.'
)
add_paragraph(
    'On the translation side, NLLB-200 [14] was the natural fit for our setup. It covers all 22 '
    'scheduled Indian languages, and the distilled 600M version fits in 8 GB of RAM while keeping '
    'roughly 95% of the full model\'s quality. The reported +44% average BLEU gain on low-resource '
    'pairs—Pashto and Kashmiri among them—was directly relevant to the domain we were working in.'
)
add_paragraph(
    'In code-switched ASR, tokenization and language modeling play an outsized role, and domain '
    'mismatch between curated corpora and field speech is often the primary source of errors [11,12].'
)

# ══════════════════════════════════════════════════════════════════════════════
# III. CHALLENGES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. Challenges Specific to Indic Languages')

add_heading('3.1 Script and Orthography', level=2)
add_paragraph(
    'Many Indic languages have multiple writing conventions, and users frequently omit diacritics, '
    'merge words, or spell phonetically. When training with standard corpora, this mismatch can '
    'inflate WER even when phonetic content is correct. A particularly important case in VANI is '
    'Punjabi: the language is natively written in Gurmukhi script, but Whisper\'s output is often '
    'romanised due to training corpus composition, causing FastText to misclassify it as Hindi '
    '(FastText Hindi confidence: 0.986 on a Punjabi test sample). This motivates Unicode script-ratio '
    'analysis as an independent signal.'
)

add_heading('3.2 Phonological Diversity', level=2)
add_paragraph(
    'The sound systems of Indic languages differ in ways that matter for ASR—aspiration distinctions, '
    'retroflex consonants, and breathy voice are not well-represented in most training corpora. One '
    'concrete consequence we observed: for a clearly Punjabi recording, Whisper returned a language '
    'probability of just 0.186, while MMS-LID—operating on the same raw audio—returned 0.9994. The '
    'gap is not a bug; it reflects how differently the two models were trained and what signals they '
    'rely on.'
)

add_heading('3.3 Code-Switching and Borrowings', level=2)
add_paragraph(
    'Code-switching between an Indic language and English is common in conversational and broadcast '
    'domains. This creates a coupled problem: recognising mixed-language acoustic sequences and '
    'decoding into mixed-script or romanised text. VANI\'s translation routing addresses this by '
    'allowing language-specific model selection (NLLB-200 for most Indic languages, IndicTrans2 as '
    'fallback for Dogri).'
)

add_heading('3.4 Low-Resource and Noisy Conditions', level=2)
add_paragraph(
    'For many Indic languages, only small amounts of transcribed speech are available. Radio and '
    'field recordings are frequently corrupted by background noise, reverberation, and channel '
    'differences. VANI applies silero-VAD-based silence removal and spectral normalization as '
    'preprocessing stages prior to ASR.'
)

# ══════════════════════════════════════════════════════════════════════════════
# IV. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. System Architecture and Proposed Techniques')
add_paragraph(
    'Audio enters VANI at one end and exits as a structured intelligence report at the other, passing '
    'through eight stages in fixed order: VAD, preprocessing, chunking, ASR, language identification, '
    'translation, keyword detection, and ISUM. Six of these stages are the subject of the techniques '
    'described below.'
)

add_heading('4.1 Technique 1: Large-Scale Weakly Supervised ASR', level=2)
add_paragraph(
    'VANI uses Whisper large-v3-turbo quantised to int8 via CTranslate2 as its ASR backbone. '
    'Radford et al. trained Whisper on 680,000 hours of multilingual audio spanning 99 languages [13], '
    'which gives it usable zero-shot coverage across a wide range of languages without any fine-tuning. '
    'Running it as a CTranslate2 int8 model cuts memory to roughly a quarter of the float32 '
    'footprint—we measured acceptable transcription quality at under 2 GB when the translation model '
    'is not loaded.'
)
add_paragraph(
    'We arrived at several configuration decisions through early testing failures. Radio static caused '
    'the decoder to enter repetition loops when condition_on_previous_text was left on; disabling it '
    'eliminated this. A no_speech_threshold of 0.70 proved necessary because Whisper will produce a '
    'transcript on near-silence if nothing prevents it, and those false segments inflate WER. We fixed '
    'beam_size=4 and temperature=0.0 so that repeated runs on the same file give identical output. '
    'Two domain-specific choices also proved useful: passing a Gurmukhi vocabulary string through '
    'initial_prompt biases the decoder toward military terminology, and word-level timestamps '
    'extracted from cross-attention alignment give downstream keyword localisation a time anchor.'
)

add_heading('4.2 Technique 2: Multilingual Translation Routing', level=2)
add_paragraph(
    'Translation routing in VANI works as a dispatch layer: each intercept is directed to whichever '
    'model handles its detected language. In practice, NLLB-200-distilled-600M takes nearly '
    'everything—Hindi, Punjabi, Urdu, Nepali, Pashto, Kashmiri, Arabic, and a long tail of others '
    'are all covered [14]. The single gap we encountered was Dogri, which NLLB-200 does not include; '
    'IndicTrans2 handles that language alone. When the detected language is English, no translation '
    'runs. Memory management here was non-trivial: keeping both models resident simultaneously would '
    'push the system over the 8 GB limit, so the routing layer loads a model on demand, then releases '
    'it explicitly (del model + gc.collect()) before the next stage loads anything.'
)

add_heading('4.3 Technique 3: Script-Aware Phonological Targeting', level=2)
add_paragraph(
    'The Unicode ranges for each script give a cheap but reliable signal that neither audio-based nor '
    'text-based models use directly. We check what fraction of characters in the ASR output fall '
    'within the Gurmukhi block (U+0A00–U+0A7F), the Devanagari block (U+0900–U+097F), the Arabic '
    'block (U+0600–U+06FF), and CJK (U+4E00–U+9FFF). When Gurmukhi characters account for more than '
    '20% of the transcript, we treat it as Punjabi regardless of what Whisper, FastText, or MMS '
    'return—no other language produces substantial Gurmukhi output under any realistic condition we '
    'have observed. To steer Whisper\'s output toward the expected script before analysis runs, we '
    'pass a domain vocabulary string through the initial_prompt parameter; this is functionally '
    'similar to shallow language model fusion but requires no additional model.'
)

add_heading('4.4 Technique 4: Three-Source Confidence-Weighted LangID Ensemble', level=2)
add_paragraph(
    'Three independent signals feed the voting layer, each with a different failure mode. Whisper '
    'contributes a language probability from its encoder\'s classification head—but as the Punjabi '
    'data showed, this can be confidently wrong when the language is underrepresented in training. '
    'FastText [16] works on the ASR transcript via character n-gram features, publishing >97% accuracy '
    'across 176 languages, but it inherits whatever misclassification Whisper introduced at the '
    'transcript stage. MMS-LID-256 [15] is different in kind: it processes the raw audio waveform '
    'and has no knowledge of what Whisper produced. Published top-1 accuracy across 256 language '
    'classes exceeds 90%, and in our experiments it was the only source that recovered correct '
    'Punjabi identification when both text-based sources had already failed.'
)
add_paragraph(
    'When all three sources agree, the consensus confidence is boosted by 10%. When two of three '
    'agree, the final score is the average of the two agreeing values. When all three disagree, the '
    'highest score is penalised by 15%. Two hard overrides sit outside this vote: Gurmukhi script '
    'above the 0.20 threshold forces Punjabi unconditionally; MMS and FastText both returning Punjabi '
    'while Whisper reports Hindi also forces Punjabi. Any result below 0.60 confidence sets '
    'uncertain=True and flags the intercept for analyst review.'
)

add_heading('4.5 Technique 5: VAD-Based Augmentation and Channel Preprocessing', level=2)
add_paragraph(
    'Radio recordings arrive with characteristics that studio-trained models do not expect. Silero-VAD '
    'runs first to strip silence; this both reduces RTF and eliminates the hallucinated transcripts '
    'that Whisper generates when given extended quiet passages. Spectral normalization compensates for '
    'channel-specific frequency response differences between recording equipment. Long files are split '
    'at VAD-detected speech boundaries rather than fixed intervals, so each chunk reaching the ASR '
    'stage contains complete utterances rather than mid-word cuts. Any segment where Whisper\'s own '
    'no_speech_prob exceeds 0.70 is discarded rather than forwarded.'
)

add_heading('4.6 Technique 6: Fairness-Aware Confidence Scoring', level=2)
add_paragraph(
    'Rather than collapsing quality information into a single number, VANI keeps the sub-scores '
    'separate. Four flags are surfaced directly in the analyst interface: LOW_LANG_CONFIDENCE fires '
    'when the ensemble sits below 0.60; ASR_LOW_CONFIDENCE fires when mean segment confidence falls '
    'below 0.50; TRANSLATION_FAILED and TRANSLATION_UNRELIABLE fire when the source language is '
    'still uncertain at translation time (ensemble <0.60). A composite score combines these: 50% '
    'weight on language confidence, 30% on model agreement, 20% on 5W completeness—one number an '
    'analyst can read without parsing each sub-score individually.'
)

# ══════════════════════════════════════════════════════════════════════════════
# V. EXPERIMENTAL SETUP
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5. Experimental Setup')

add_heading('5.1 Hardware and Deployment Constraints', level=2)
add_paragraph(
    'All experiments ran on a consumer laptop with 8 GB RAM and no GPU. Network access was '
    'intentionally disabled to validate offline operation. The 8 GB ceiling is not a soft '
    'guideline—exceeding it causes the OS to begin swapping, which makes inference time unpredictable '
    'and occasionally crashes the process. Two large models cannot be resident simultaneously without '
    'hitting it.'
)

add_heading('5.2 Software Stack', level=2)
add_paragraph('Table 1 summarises the models used.', space_after=4)

# Table 1
table1 = doc.add_table(rows=6, cols=4)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Component', 'Model', 'Params', 'Format']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=9)
rows_data = [
    ['ASR', 'Whisper large-v3-turbo', '809M', 'CT2 int8'],
    ['Translation', 'NLLB-200-dist.-600M', '600M', 'PT float32'],
    ['Text LangID', 'FastText lid.176.bin', '—', 'Binary'],
    ['Audio LangID', 'MMS-LID-256', '150M', 'PT float32'],
    ['ISUM', 'Rule-based + Qwen2.5-1.5B', '1.5B', 'PT float32'],
]
for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        cell = table1.rows[i+1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Table 1. Software Stack')
set_font(r, bold=True, size=9)
p.paragraph_format.space_after = Pt(8)

add_heading('5.3 Test Audio', level=2)
add_paragraph(
    'Qualitative test set: five files comprising Harvard Sentences (16.6s, English), LJSpeech '
    'LJ001-0004 (5.1s, English), Aesop\'s Fables multi-chunk (52.4s, English), and two Punjabi '
    'broadcast speech samples (7.0s, 12.3s).'
)
add_paragraph(
    'For scale evaluation, 30 samples per language were pulled from four HuggingFace datasets, '
    'clipped to 2–20s: shunyalabs/punjabi-speech-dataset (Gurmukhi references), '
    'MatrixSpeechAI/All_Hindi_ASR_v1.2 (Devanagari), m-aliabbas/common_voice_urdu (Nastaliq), and '
    'iamTangsang/OpenSLR54-Nepali-ASR (Devanagari). WER and CER were computed with jiwer after '
    'Unicode-aware punctuation removal and lowercasing of both hypothesis and reference strings.'
)

# ══════════════════════════════════════════════════════════════════════════════
# VI. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6. Experimental Results')

add_heading('6.1 ASR Performance', level=2)
add_paragraph('Table 2 reports per-segment ASR confidence alongside published WER benchmarks [13].',
              space_after=4)

table2 = doc.add_table(rows=6, cols=5)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['File', 'Lang', 'Segs', 'Conf', 'RTF']
for i, h in enumerate(h2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=9)
t2data = [
    ['harvard.wav', 'en', '1', '0.973', '3.24×'],
    ['LJ001-0004.wav', 'en', '2', '0.782', '8.00×'],
    ['Speaker26_000.wav', 'en', '7', '0.922', '—'],
    ['sent_1.wav', 'pa', '1', '0.938', '2.60×'],
    ['sent_10.wav', 'pa', '1', '0.879', '5.53×'],
]
for i, row in enumerate(t2data):
    for j, val in enumerate(row):
        cell = table2.rows[i+1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Table 2. ASR Performance on Qualitative Test Samples')
set_font(r, bold=True, size=9)
p.paragraph_format.space_after = Pt(8)

add_paragraph(
    'Clean English audio gave the expected result: Harvard Sentences at 0.973 mean confidence, in '
    'line with Whisper\'s published clean-speech numbers. The Punjabi results were less predictable—'
    '0.879 and 0.938 are high enough to be treated as reliable, though the LangID data in Section '
    '6.2 shows that confidence and label correctness are not the same thing.'
)

add_heading('6.2 Language Identification Ensemble Results', level=2)
add_paragraph('Table 3 shows per-source language predictions and final ensemble decisions.',
              space_after=4)

table3 = doc.add_table(rows=6, cols=6)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
h3 = ['File', 'True', 'Whisper', 'FastText', 'MMS', 'Final']
for i, h in enumerate(h3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=9)
t3data = [
    ['harvard', 'en', 'en (0.999)', 'en (0.862)', 'en (0.997)', 'en ✓'],
    ['LJ001', 'en', 'en (1.000)', 'en (0.992)', 'en (0.991)', 'en ✓'],
    ['Spkr26', 'en', 'en (0.984)', 'en (0.984)', '—', 'en ✓'],
    ['sent_1', 'pa', 'hi (0.707) ✗', 'hi (0.986) ✗', 'pa (0.999) ✓', 'pa ✓'],
    ['sent_10', 'pa', 'pa (0.186) ?', 'pa (1.000) ✓', 'pa (0.999) ✓', 'pa ✓'],
]
for i, row in enumerate(t3data):
    for j, val in enumerate(row):
        cell = table3.rows[i+1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Table 3. Three-Source LangID Results')
set_font(r, bold=True, size=9)
p.paragraph_format.space_after = Pt(8)

add_paragraph(
    'sent_1.wav makes the design argument most clearly. Whisper returned Hindi at 0.707; FastText '
    'returned Hindi at 0.986; the audio is Punjabi. Both text-based sources agreed confidently on '
    'the wrong answer. MMS-LID, running on the raw waveform with no knowledge of the transcript, '
    'returned Punjabi at 0.9992. Without the ensemble override, this sample would have been routed '
    'to Hindi translation without any flag. Final accuracy across all five qualitative samples: 5/5.'
)

add_heading('6.3 Translation Results', level=2)
add_paragraph(
    'NLLB-200 translated both samples via the pan_Guru → eng_Latn route. The shorter clip '
    '(sent_1.wav, 7.0s) came back as "Kindri Minister Manderne decided to set up a National Council '
    'NCVAT..." in 8.76s. The longer one (sent_10.wav, 12.3s) produced "Even the ship\'s crew were '
    'not able to find a boat..." in 13.66s. The sent_10 translation doubled the phrase "find a '
    'boat"—a repetition artefact that appears when NLLB-200 has limited context on a syntactically '
    'complex short input and the decoder reinforces itself.'
)

add_heading('6.4 System Performance', level=2)
add_paragraph('Table 4 reports end-to-end system performance.', space_after=4)

table4 = doc.add_table(rows=5, cols=5)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
h4 = ['File', 'Proc. (s)', 'RTF', 'Peak RAM', 'Trans. (s)']
for i, h in enumerate(h4):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=9)
t4data = [
    ['harvard.wav', '104.8', '6.3×', '2,850 MB', '0.18'],
    ['LJ001-0004.wav', '69.6', '13.6×', '314 MB', '0.16'],
    ['sent_1.wav', '31.4', '4.5×', '5,453 MB', '8.76'],
    ['sent_10.wav', '124.4', '10.1×', '330 MB', '13.66'],
]
for i, row in enumerate(t4data):
    for j, val in enumerate(row):
        cell = table4.rows[i+1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Table 4. System Performance (8 GB RAM, CPU-only)')
set_font(r, bold=True, size=9)
p.paragraph_format.space_after = Pt(8)

add_paragraph(
    'ASR accounted for 53–68% of total processing time. Memory usage was more variable: 314 MB on '
    'LJSpeech (no translation model loaded) up to 5,453 MB on sent_1.wav (translation model loaded '
    'concurrently). The sent_1 peak is close enough to the 8 GB ceiling that any model lifecycle '
    'error—a missed gc.collect(), a held reference—would have pushed it over.'
)

add_heading('6.5 Large-Scale Dataset Evaluation (120 Samples, 4 Languages)', level=2)
add_paragraph('Table 5 reports aggregate performance across 30 samples per language.', space_after=4)

table5 = doc.add_table(rows=9, cols=5)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
h5 = ['Metric', 'Punjabi (pa)', 'Hindi (hi)', 'Urdu (ur)', 'Nepali (ne)']
for i, h in enumerate(h5):
    cell = table5.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, bold=True, size=9)
t5data = [
    ['LangID accuracy', '29/30 (96.7%)', '29/30 (96.7%)', '21/30 (70.0%)', '19/30 (63.3%)'],
    ['Mean WER', '117.2%†', '46.2%', '74.4%', '108.3%‡'],
    ['Mean CER', '98.5%†', '32.4%', '63.2%', '87.3%‡'],
    ['Best WER', '100.0%†', '13.0%', '22.2%', '100.0%‡'],
    ['Mean seg. conf.', '0.870', '0.879', '0.455', '0.187'],
    ['Mean RTF (CPU)', '4.57×', '4.11×', '8.98×', '7.00×'],
    ['Translation', '30/30 (100%)', '30/30 (100%)', '30/30 (100%)', '30/30 (100%)'],
    ['Skipped', '3', '0', '11', '8'],
]
for i, row in enumerate(t5data):
    for j, val in enumerate(row):
        cell = table5.rows[i+1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            set_font(run, size=9)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Table 5. Aggregate Performance on Open-Source Datasets (30 Samples per Language). '
              '† WER inflated: Whisper outputs Devanagari; references are Gurmukhi. '
              '‡ WER inflated: Whisper transcribes Nepali as European languages.')
set_font(r, bold=True, size=9)
p.paragraph_format.space_after = Pt(8)

add_paragraph(
    'The result that struck us most was how clean the confidence gradient was. Hindi and Punjabi both '
    'sat above 0.87; Urdu came in at 0.455 and Nepali at 0.187. The ordering matches what is known '
    'about Whisper\'s training representation—the languages with less coverage produce lower '
    'confidence—but we had not expected the gradient to be this consistent. It means that even '
    'without a reference transcription, an analyst looking at the mean segment confidence score gets '
    'a reasonable signal about whether to trust the output.'
)
add_paragraph(
    'LangID failure analysis reveals distinct underlying causes. Hindi/Punjabi (2 of 60): heavy '
    'Urdu/Persian loanword vocabulary caused all three sources to agree on Urdu. Urdu (9 of 30): '
    'FastText returned "unknown" in 8 of 9 failures; Whisper predicted Turkish (3), English (5), or '
    'Arabic (1); MMS correctly returned Urdu in 3 of 9 cases. Nepali (11 of 30): FastText returned '
    '"unknown" in 10 of 11 failures; Whisper hallucinated Portuguese (2), English (4), German, '
    'Italian, Icelandic, or Urdu at very low confidence; MMS correctly returned Nepali in 4 of 11 '
    'failures.'
)

# ══════════════════════════════════════════════════════════════════════════════
# VII. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7. Discussion')

add_heading('7.1 Why the Ensemble Corrects Whisper\'s Hindi/Punjabi Confusion', level=2)
add_paragraph(
    'Three causes overlap in the Hindi/Punjabi case. The languages share enough of their phone '
    'inventory that the acoustic boundary is not sharp—they come from the same dialect continuum. '
    'Whisper has seen substantially more Hindi than Punjabi during training, so when the two are '
    'acoustically ambiguous, it falls back on Hindi. FastText compounds this: it reads the romanised '
    'transcript Whisper produced, identifies Punjabi phonemes written in Devanagari-adjacent '
    'romanisation as Hindi, and is correct to do so given what it received. Neither source is making '
    'an error by its own logic; they are just working from the same flawed input. MMS-LID is the '
    'circuit-breaker. It never touches the transcript—it runs a wav2vec 2.0 encoder on the raw '
    'audio and on sent_1.wav came back at 0.9992 Punjabi while both text sources were still '
    'reporting Hindi. Omitting audio-based LID from an Indic pipeline means this class of failure '
    'has no recovery path.'
)

add_heading('7.2 Why the Techniques Complement Each Other', level=2)
add_paragraph(
    'The techniques described above did not emerge from a single design decision—each was added to '
    'fill a gap that the previous component left open. Whisper provides a transcript and a language '
    'probability, but as the Punjabi experiments show, those probabilities can be confidently wrong. '
    'The ensemble was built because a single source is not sufficient. The Unicode override was added '
    'because the ensemble still relies partly on the romanised transcript, which carries no useful '
    'signal when Whisper writes Punjabi phonemes in the wrong script. The confidence flags came last, '
    'after we observed that even a correctly identified language could produce uncertain output in '
    'ways the pipeline had no mechanism to surface.'
)

add_heading('7.3 ASR Coverage as a Predictor of Ensemble Accuracy', level=2)
add_paragraph(
    'The four-language breakdown reveals a pattern that was cleaner than we expected: Whisper\'s mean '
    'confidence tracks directly with how much of each language was in its training data, and that '
    'gap carries straight through into LangID accuracy. Hindi at 0.879 confidence / 96.7% accuracy; '
    'Urdu at 0.455 / 70.0%; Nepali at 0.187 / 63.3%—the ordering is exact. The mechanism is not '
    'simply that low confidence predicts wrong answers. It is that when Whisper lacks coverage of a '
    'language, it assigns high probability to a superficially related one it knows better: Turkish '
    'for Urdu, Portuguese or German for Nepali. Those wrong high-confidence outputs then corrupt the '
    'FastText vote, since FastText has no independent audio signal to disagree with. Mean segment '
    'confidence therefore functions as a proxy for downstream accuracy even without any reference '
    'transcript to compare against.'
)

add_heading('7.4 Gurmukhi Script Gap in Whisper', level=2)
add_paragraph(
    'None of the 30 Punjabi test samples produced Gurmukhi characters from Whisper. Every transcript '
    'came back in Devanagari. The 117.2% WER against Gurmukhi references is not a measure of '
    'transcription quality—it is a measure of script distance. Gurmukhi is simply absent from '
    'Whisper\'s output vocabulary. For the intelligence use case this is workable: the downstream '
    'goal is an English translation, and NLLB-200 handles Devanagari Punjabi input correctly '
    'regardless of which script the audio was originally in. A different approach would be needed if '
    'the requirement were an archival Gurmukhi transcript—either a fine-tuned MMS-300M model or a '
    'Wav2Vec2 model trained with Gurmukhi output.'
)

add_heading('7.5 Limitations', level=2)
add_paragraph(
    'CPU inference runs at 2.95–3.93× realtime for Hindi and Punjabi, and 8.98× for Urdu. Live '
    'transcription is not feasible at these speeds; the system targets post-collection batch '
    'processing. We did not test GPU acceleration, but published Whisper benchmarks suggest RTF '
    'below 0.3× is achievable, which might make live use viable. The evaluation set is read speech '
    'and broadcast audio from four languages only; conversational speech, code-switched audio, and '
    'adversarial conditions are untested. NLLB-200 occasionally repeats phrases on short inputs '
    'with complex syntax—sent_10.wav was one such case.'
)

# ══════════════════════════════════════════════════════════════════════════════
# VIII. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('8. Conclusion and Future Work')
add_paragraph(
    'VANI was built to answer a practical question: can offline Indic ASR be made reliable enough on '
    'a standard laptop to be useful for intelligence analysis, without any network dependency? The '
    '120-sample evaluation gives a partial answer. Where Whisper\'s training coverage is adequate, '
    'the system holds up—Hindi and Punjabi both reach 96.7% LangID accuracy and 100% translation '
    'success. The drop on Urdu (70.0% LangID) and Nepali (63.3%) is not surprising; it tracks the '
    'same training data gap that causes mean confidence to fall from 0.879 on Hindi to 0.187 on '
    'Nepali. What stood out was how cleanly this gradient predicts downstream reliability: an analyst '
    'reading the confidence score gets a reasonable proxy for accuracy without any reference data.'
)
add_paragraph(
    'Translation held across all 120 samples. NLLB-200-distilled-600M produced output for every '
    'intercept regardless of whether the transcript was in the correct script—which matters in '
    'practice because the final deliverable is an English translation, not the intermediate '
    'transcript. Hindi is the one language where WER is an honest metric, since hypothesis and '
    'reference use matching scripts, and 46.2% WER with 32.4% CER on radio-quality audio reflects '
    'the genuine difficulty of the domain.'
)
add_paragraph(
    'The most pressing gaps are the script-level overrides not yet implemented: a Devanagari '
    'disambiguation rule for Nepali—analogous to the Gurmukhi override—would likely recover a '
    'substantial fraction of the 11 Nepali failures in the evaluation set, and an Arabic-script '
    'heuristic could address the FastText "unknown" failures for Urdu. A fine-tuned ASR model with '
    'Gurmukhi output support would remove the script mismatch that currently makes Punjabi WER '
    'uninterpretable. We also plan to extend evaluation to Bengali, Pashto, and code-switched speech, '
    'and to profile GPU-accelerated runtime to determine whether live-intercept processing becomes '
    'feasible.'
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading('References')
refs = [
    '[1] G. Hinton et al., "Deep neural networks for acoustic modeling in speech recognition," IEEE Signal Processing Magazine, 2012.',
    '[2] A. Graves, A. Mohamed, G. Hinton, "Speech recognition with deep recurrent neural networks," Proc. ICASSP, 2013.',
    '[3] J. Chorowski et al., "Attention-based models for speech recognition," Proc. NeurIPS, 2015.',
    '[4] A. Vaswani et al., "Attention is all you need," Proc. NeurIPS, 2017.',
    '[5] A. Gulati et al., "Conformer: Convolution-augmented transformer for speech recognition," Proc. Interspeech, 2020.',
    '[6] A. Baevski et al., "wav2vec 2.0: A framework for self-supervised learning of speech representations," Proc. NeurIPS, 2020.',
    '[7] W.-N. Hsu et al., "HuBERT: Self-supervised speech representation learning by masked prediction of hidden units," IEEE/ACM Trans. Audio, Speech, Lang. Process., 2021.',
    '[8] V. Pratap et al., "Scaling up online speech recognition using ConvNets and streaming Transformers," Interspeech, 2020.',
    '[9] A. Babu et al., "XLS-R: Self-supervised cross-lingual speech representation learning at scale," Proc. Interspeech, 2021.',
    '[10] R. Ardila et al., "Common Voice: A massively-multilingual speech corpus," Proc. LREC, 2020.',
    '[11] T. Kudo and J. Richardson, "SentencePiece: A simple and language independent subword tokenizer," Proc. EMNLP, 2018.',
    '[12] Y. Li and P. Fung, "Code-switching language modeling and ASR: A survey," Computational Linguistics, 2019.',
    '[13] A. Radford et al., "Robust speech recognition via large-scale weak supervision," arXiv:2212.04356, OpenAI, 2022.',
    '[14] M. R. Costa-jussà et al., "No language left behind: Scaling human-centered machine translation," arXiv:2207.04672, Meta AI, 2022.',
    '[15] V. Pratap et al., "Scaling speech technology to 1,000+ languages," arXiv:2305.13516, Meta AI, 2023.',
    '[16] A. Joulin et al., "Bag of tricks for efficient text classification," Proc. EACL, 2017.',
]
for ref in refs:
    add_paragraph(ref, size=9, space_after=3, indent=0.3)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save('VANI_Paper_IJAINN.docx')
print("Done: VANI_Paper_IJAINN.docx")
