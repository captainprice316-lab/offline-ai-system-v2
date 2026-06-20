"""
src/report_exporter.py – VANI Report Exporter
----------------------------------------------
Generates formatted ISUM reports as DOCX or PDF from a pipeline result dict.
"""

import io
from datetime import datetime

# ── DOCX ───────────────────────────────────────────────────────────────────────

def build_docx(result: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    isum = result.get("isum", {})
    doc  = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Header ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("VANI – INTELLIGENCE SUMMARY REPORT", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, 0x00, 0x99, 0x55)

    doc.add_paragraph()  # spacer

    # ── Report metadata ────────────────────────────────────────────────────────
    _add_field(doc, "REPORT ID",    isum.get("report_id", "-"))
    _add_field(doc, "TIMESTAMP",    isum.get("timestamp_utc", "-"))
    _add_field(doc, "AUDIO FILE",   isum.get("audio_file", "-"))
    _add_field(doc, "PROC TIME",    f"{isum.get('processing_time_s', 0)}s")
    _add_field(doc, "LANGUAGE",     (result.get("final_language") or "-").upper())
    _add_field(doc, "THREAT LEVEL", isum.get("threat_level", "CLEAR"))

    doc.add_paragraph()

    # ── Assessment ─────────────────────────────────────────────────────────────
    _section_heading(doc, "ASSESSMENT")
    doc.add_paragraph(isum.get("assessment", "-"))

    # ── 5W fields ──────────────────────────────────────────────────────────────
    _section_heading(doc, "FIVE-W INTELLIGENCE FIELDS")
    _add_field(doc, "WHO  – Actors",    isum.get("who",   "Not identified."))
    _add_field(doc, "WHAT – Activity",  isum.get("what",  "No activity detected."))
    _add_field(doc, "WHERE – Location", isum.get("where", "Not identified."))
    _add_field(doc, "WHEN – Temporal",  isum.get("when",  "No temporal reference."))

    # ── Keyword categories ─────────────────────────────────────────────────────
    cats = isum.get("top_categories", [])
    if cats:
        _section_heading(doc, "TRIGGERED CATEGORIES")
        doc.add_paragraph("  •  " + "\n  •  ".join(cats))

    # ── Quality flags ──────────────────────────────────────────────────────────
    flags = isum.get("confidence_flags", [])
    _section_heading(doc, "QUALITY FLAGS")
    doc.add_paragraph("  •  ".join(flags) if flags else "NO FLAGS – HIGH CONFIDENCE RESULT")

    # ── Speaker-labelled transcript ────────────────────────────────────────────
    _segs = result.get("segments", [])
    _spk_tx = ""
    if _segs and _segs[0].get("speaker"):
        try:
            from diarize_module import build_speaker_transcript, speaker_stats
            _spk_tx    = build_speaker_transcript(_segs)
            _spk_stats = speaker_stats(_segs)
        except Exception:
            _spk_tx    = ""
            _spk_stats = {}
    else:
        _spk_stats = {}

    lang_code  = (result.get("final_language") or "?").upper()
    lang_names = {
        "HI":"HINDI","PA":"PUNJABI","UR":"URDU","NE":"NEPALI","DOI":"DOGRI",
        "PS":"PASHTO","ZH":"MANDARIN","MY":"BURMESE","KS":"KASHMIRI",
        "EN":"ENGLISH","MAI":"MAITHILI","BN":"BENGALI","BO":"TIBETAN",
    }
    lang_name = lang_names.get(lang_code, lang_code)

    full_transcript = result.get("transcript", "") or isum.get("transcript_snippet", "")

    if _spk_tx:
        n_spk = len(_spk_stats)
        _section_heading(doc, f"Speaker-Labelled Transcript  [{lang_name} · {lang_code}]  —  {n_spk} speaker(s)")
        for line in _spk_tx.splitlines():
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)
        if _spk_stats:
            doc.add_paragraph()
            _add_field(doc, "Speaker breakdown",
                       "  ·  ".join(f"{s}: {v['words']} words / {v['segments']} segs"
                                    for s, v in sorted(_spk_stats.items())))
    elif full_transcript:
        _section_heading(doc, f"Intercept Transcript  [{lang_name} · {lang_code}]")
        p = doc.add_paragraph(full_transcript)
        p.style.font.size = Pt(9)

    # ── Full English translation ────────────────────────────────────────────────
    trans_obj  = result.get("translation", {})
    full_trans = (trans_obj.get("translated_text", "")
                  if isinstance(trans_obj, dict) else str(trans_obj))
    if full_trans and full_trans != full_transcript and lang_code != "EN":
        _section_heading(doc, "English Translation  [EN]")
        p = doc.add_paragraph(full_trans)
        p.style.font.size = Pt(9)
    elif lang_code == "EN":
        _section_heading(doc, "English Translation  [EN]")
        p = doc.add_paragraph("Source language is English — no translation required.")
        p.style.font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ── Preprocessing / SNR ────────────────────────────────────────────────────
    _pre = result.get("preprocessing", {}) or {}
    if _pre.get("snr_before_db") is not None:
        _section_heading(doc, "Audio Preprocessing")
        _snr_b = _pre["snr_before_db"]
        _snr_a = _pre.get("snr_after_db", _snr_b)
        _add_field(doc, "SNR (before denoising)", f"{_snr_b:.1f} dB")
        _add_field(doc, "SNR (after denoising)",  f"{_snr_a:.1f} dB  ({'+' if _snr_a >= _snr_b else ''}{_snr_a - _snr_b:.1f} dB)")
        _add_field(doc, "Duration after processing", f"{_pre.get('duration_sec', '?')}s")

    # ── Footer ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Generated by VANI Offline Intelligence System  |  {datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size  = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_field(doc, label: str, value: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run_lbl = p.add_run(f"{label}: ")
    run_lbl.bold = True
    run_lbl.font.size = Pt(10)
    run_val = p.add_run(str(value))
    run_val.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def _section_heading(doc, text: str):
    from docx.shared import Pt, RGBColor
    h = doc.add_heading(text, level=2)
    _set_heading_color(h, 0x00, 0x88, 0x44)


def _set_heading_color(heading, r, g, b):
    from docx.shared import RGBColor
    for run in heading.runs:
        run.font.color.rgb = RGBColor(r, g, b)


# ── PDF ────────────────────────────────────────────────────────────────────────

def build_pdf(result: dict, metrics: dict = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
    )
    from reportlab.platypus.flowables import HRFlowable

    isum  = result.get("isum", {})
    buf   = io.BytesIO()
    doc   = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2*cm, bottomMargin=2.2*cm,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
    )

    W = A4[0] - 4.4*cm   # usable width

    # ── Colour palette (light/professional theme) ──────────────────────────────
    WHITE     = colors.white
    PAGE_BG   = colors.white
    NAVY      = colors.HexColor("#0d1b2a")
    DARK      = colors.HexColor("#1a2535")
    MID       = colors.HexColor("#3a4f63")
    MUTED     = colors.HexColor("#5a7080")
    RULE      = colors.HexColor("#c8d8e4")
    ROW_ALT   = colors.HexColor("#f4f7fa")
    HDR_BG    = colors.HexColor("#0d1b2a")
    GREEN     = colors.HexColor("#007744")
    GREEN_LT  = colors.HexColor("#e6f4ee")

    THREAT_BG = {
        "CRITICAL": colors.HexColor("#ffebee"),
        "HIGH":     colors.HexColor("#fff3e0"),
        "MEDIUM":   colors.HexColor("#fffde7"),
        "LOW":      colors.HexColor("#e3f2fd"),
        "CLEAR":    colors.HexColor("#e8f5e9"),
    }
    THREAT_FG = {
        "CRITICAL": colors.HexColor("#c62828"),
        "HIGH":     colors.HexColor("#e65100"),
        "MEDIUM":   colors.HexColor("#f57f17"),
        "LOW":      colors.HexColor("#1565c0"),
        "CLEAR":    colors.HexColor("#2e7d32"),
    }
    threat     = isum.get("threat_level", "CLEAR")
    t_bg       = THREAT_BG.get(threat, GREEN_LT)
    t_fg       = THREAT_FG.get(threat, GREEN)

    # ── Styles ─────────────────────────────────────────────────────────────────
    title_style = ParagraphStyle(
        "VTitle", fontSize=22, textColor=WHITE,
        fontName="Helvetica-Bold", alignment=1,
        leading=26, spaceAfter=0,
    )
    subtitle_style = ParagraphStyle(
        "VSub", fontSize=9, textColor=colors.HexColor("#a0b8c8"),
        fontName="Helvetica", alignment=1,
        leading=13, spaceAfter=0,
    )
    section_style = ParagraphStyle(
        "VSection", fontSize=7, textColor=GREEN,
        fontName="Helvetica-Bold", spaceBefore=14,
        spaceAfter=3, leading=10, uppercase=1,
    )
    body_style = ParagraphStyle(
        "VBody", fontSize=10, textColor=DARK,
        fontName="Helvetica", spaceAfter=4,
        leading=15,
    )
    label_style = ParagraphStyle(
        "VLabel", fontSize=8, textColor=MUTED,
        fontName="Helvetica-Bold", spaceAfter=0,
        leading=11,
    )
    mono_style = ParagraphStyle(
        "VMono", fontSize=8, textColor=MID,
        fontName="Courier", spaceAfter=4,
        leading=12, backColor=ROW_ALT,
        borderPad=4,
    )
    badge_style = ParagraphStyle(
        "VBadge", fontSize=11, textColor=t_fg,
        fontName="Helvetica-Bold", alignment=1,
        leading=14, backColor=t_bg,
        borderPad=5, spaceAfter=0,
    )
    footer_style = ParagraphStyle(
        "VFooter", fontSize=7, textColor=MUTED,
        fontName="Helvetica", alignment=1, leading=10,
    )

    story = []

    # ── Title banner ───────────────────────────────────────────────────────────
    banner = Table(
        [[Paragraph("VANI", title_style)],
         [Paragraph("VOICE ANALYSIS &amp; NEURAL INTELLIGENCE", subtitle_style)],
         [Paragraph("INTELLIGENCE SUMMARY REPORT", subtitle_style)]],
        colWidths=[W],
    )
    banner.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), HDR_BG),
        ("TOPPADDING",    (0,0), (-1,-1), 10),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ("LEFTPADDING",   (0,0), (-1,-1), 0),
        ("RIGHTPADDING",  (0,0), (-1,-1), 0),
    ]))
    story.append(banner)
    story.append(Spacer(1, 0.3*cm))

    # ── Threat badge ───────────────────────────────────────────────────────────
    badge_tbl = Table([[Paragraph(f"THREAT LEVEL: {threat}", badge_style)]], colWidths=[W])
    badge_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), t_bg),
        ("TOPPADDING",    (0,0), (-1,-1), 7),
        ("BOTTOMPADDING", (0,0), (-1,-1), 7),
        ("LEFTPADDING",   (0,0), (-1,-1), 8),
        ("BOX",           (0,0), (-1,-1), 1, t_fg),
    ]))
    story.append(badge_tbl)
    story.append(Spacer(1, 0.4*cm))

    # ── Metadata table ─────────────────────────────────────────────────────────
    meta_rows = [
        [_lv("REPORT ID",  label_style),  _lv(isum.get("report_id","-"),  body_style),
         _lv("LANGUAGE",   label_style),  _lv((result.get("final_language") or "-").upper(), body_style)],
        [_lv("TIMESTAMP",  label_style),  _lv(isum.get("timestamp_utc","-"), body_style),
         _lv("PROC TIME",  label_style),  _lv(f"{isum.get('processing_time_s',0)}s", body_style)],
        [_lv("AUDIO FILE", label_style),  _lv(isum.get("audio_file","-"), body_style),
         _lv("ROUTE",      label_style),  _lv(result.get("translation_route","-"), body_style)],
    ]
    meta_tbl = Table(meta_rows, colWidths=[2.8*cm, 6*cm, 2.8*cm, W-11.6*cm])
    meta_tbl.setStyle(TableStyle([
        ("GRID",           (0,0), (-1,-1), 0.5, RULE),
        ("ROWBACKGROUNDS", (0,0), (-1,-1), [WHITE, ROW_ALT]),
        ("VALIGN",         (0,0), (-1,-1), "TOP"),
        ("TOPPADDING",     (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",  (0,0), (-1,-1), 5),
        ("LEFTPADDING",    (0,0), (-1,-1), 6),
        ("RIGHTPADDING",   (0,0), (-1,-1), 6),
    ]))
    story.append(meta_tbl)

    # ── Assessment ─────────────────────────────────────────────────────────────
    story.append(_sec("ASSESSMENT", section_style))
    story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
    story.append(Spacer(1, 0.15*cm))
    assess_tbl = Table([[Paragraph(isum.get("assessment","-"), body_style)]], colWidths=[W])
    assess_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,-1), GREEN_LT),
        ("BOX",           (0,0), (-1,-1), 1, colors.HexColor("#b2dfcc")),
        ("LEFTPADDING",   (0,0), (-1,-1), 10),
        ("RIGHTPADDING",  (0,0), (-1,-1), 10),
        ("TOPPADDING",    (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
    ]))
    story.append(assess_tbl)

    # ── 5W table ───────────────────────────────────────────────────────────────
    story.append(_sec("FIVE-W INTELLIGENCE FIELDS", section_style))
    story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
    story.append(Spacer(1, 0.15*cm))

    w_hdr_style = ParagraphStyle("WHdr", fontSize=9, textColor=WHITE,
                                  fontName="Helvetica-Bold", leading=12)
    w_lbl_style = ParagraphStyle("WLbl", fontSize=8, textColor=MUTED,
                                  fontName="Helvetica-Bold", leading=11)
    w_val_style = ParagraphStyle("WVal", fontSize=9, textColor=DARK,
                                  fontName="Helvetica", leading=13)
    five_w = [
        [Paragraph("KEY", w_hdr_style), Paragraph("FIELD", w_hdr_style), Paragraph("EXTRACTED CONTENT", w_hdr_style)],
        [Paragraph("WHO",   w_lbl_style), Paragraph("Actors Identified",   w_lbl_style), Paragraph(isum.get("who",   "Not identified."), w_val_style)],
        [Paragraph("WHAT",  w_lbl_style), Paragraph("Activity Detected",   w_lbl_style), Paragraph(isum.get("what",  "No activity detected."), w_val_style)],
        [Paragraph("WHERE", w_lbl_style), Paragraph("Location Indicators", w_lbl_style), Paragraph(isum.get("where", "Not identified."), w_val_style)],
        [Paragraph("WHEN",  w_lbl_style), Paragraph("Temporal Indicators", w_lbl_style), Paragraph(isum.get("when",  "No temporal reference."), w_val_style)],
    ]
    w_tbl = Table(five_w, colWidths=[1.5*cm, 4*cm, W-5.5*cm])
    w_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
        ("TEXTCOLOR",     (0,0), (-1,0),  WHITE),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
        ("GRID",          (0,0), (-1,-1), 0.5, RULE),
        ("VALIGN",        (0,0), (-1,-1), "TOP"),
        ("TOPPADDING",    (0,0), (-1,-1), 6),
        ("BOTTOMPADDING", (0,0), (-1,-1), 6),
        ("LEFTPADDING",   (0,0), (-1,-1), 6),
        ("RIGHTPADDING",  (0,0), (-1,-1), 6),
    ]))
    story.append(w_tbl)

    # ── Categories ─────────────────────────────────────────────────────────────
    cats = isum.get("top_categories", [])
    if cats:
        story.append(_sec("TRIGGERED CATEGORIES", section_style))
        story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
        story.append(Spacer(1, 0.15*cm))
        cat_style = ParagraphStyle("Cat", fontSize=9, textColor=DARK,
                                    fontName="Helvetica", leading=14)
        story.append(Paragraph("  ·  ".join(c.replace("_"," ").upper() for c in cats), cat_style))

    # ── Quality flags ──────────────────────────────────────────────────────────
    flags = isum.get("confidence_flags", [])
    story.append(_sec("QUALITY FLAGS", section_style))
    story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
    story.append(Spacer(1, 0.15*cm))
    if flags:
        flag_style = ParagraphStyle("Flag", fontSize=9, textColor=colors.HexColor("#b45309"),
                                     fontName="Helvetica-Bold", leading=14,
                                     backColor=colors.HexColor("#fffbeb"), borderPad=4)
        story.append(Paragraph("  ·  ".join(flags), flag_style))
    else:
        ok_style = ParagraphStyle("FlagOK", fontSize=9, textColor=GREEN,
                                   fontName="Helvetica-Bold", leading=14,
                                   backColor=GREEN_LT, borderPad=4)
        story.append(Paragraph("NO FLAGS – HIGH CONFIDENCE RESULT", ok_style))

    # ── Full transcript (detected language) + English translation ──────────────
    full_transcript = result.get("transcript", "") or isum.get("transcript_snippet", "")
    trans_obj       = result.get("translation", {})
    full_trans      = (trans_obj.get("translated_text", "")
                       if isinstance(trans_obj, dict) else str(trans_obj))
    lang_code  = (result.get("final_language") or "?").upper()
    lang_names = {
        "HI":"HINDI","PA":"PUNJABI","UR":"URDU","NE":"NEPALI","DOI":"DOGRI",
        "PS":"PASHTO","ZH":"MANDARIN","MY":"BURMESE","KS":"KASHMIRI",
        "EN":"ENGLISH","MAI":"MAITHILI","BN":"BENGALI","BO":"TIBETAN",
    }
    lang_name = lang_names.get(lang_code, lang_code)

    # Check for speaker diarization
    _segs  = result.get("segments", [])
    _spk_tx, _spk_stats = "", {}
    if _segs and _segs[0].get("speaker"):
        try:
            from diarize_module import build_speaker_transcript, speaker_stats
            _spk_tx    = build_speaker_transcript(_segs)
            _spk_stats = speaker_stats(_segs)
        except Exception:
            pass

    if full_transcript or full_trans or _spk_tx:
        n_spk = len(_spk_stats)
        if _spk_tx:
            tx_heading = f"SPEAKER-LABELLED TRANSCRIPT  [{lang_name} · {lang_code}]  —  {n_spk} SPEAKER(S)"
        else:
            tx_heading = "INTERCEPT TRANSCRIPT"
        story.append(_sec(tx_heading, section_style))
        story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
        story.append(Spacer(1, 0.15*cm))

        tx_lbl_style  = ParagraphStyle("TxLbl", fontSize=7, textColor=colors.HexColor("#1565c0"),
                                        fontName="Helvetica-Bold", leading=10, spaceAfter=3)
        tx_en_style   = ParagraphStyle("TxEn",  fontSize=7, textColor=GREEN,
                                        fontName="Helvetica-Bold", leading=10, spaceAfter=3)
        spk_hdr_style = ParagraphStyle("SpkH",  fontSize=8, textColor=WHITE,
                                        fontName="Helvetica-Bold", leading=11)
        spk_val_style = ParagraphStyle("SpkV",  fontSize=9, textColor=DARK,
                                        fontName="Helvetica", leading=13)

        # Speaker breakdown mini-table
        if _spk_stats:
            spk_rows = [[Paragraph("SPEAKER",   spk_hdr_style),
                         Paragraph("WORDS",     spk_hdr_style),
                         Paragraph("SEGMENTS",  spk_hdr_style)]]
            for spk, sv in sorted(_spk_stats.items()):
                spk_rows.append([
                    Paragraph(spk,              spk_val_style),
                    Paragraph(str(sv["words"]), spk_val_style),
                    Paragraph(str(sv["segments"]), spk_val_style),
                ])
            spk_tbl = Table(spk_rows, colWidths=[4*cm, 3*cm, 3*cm])
            spk_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
                ("TEXTCOLOR",     (0,0), (-1,0),  WHITE),
                ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
                ("GRID",          (0,0), (-1,-1), 0.5, RULE),
                ("TOPPADDING",    (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                ("LEFTPADDING",   (0,0), (-1,-1), 6),
            ]))
            story.append(spk_tbl)
            story.append(Spacer(1, 0.2*cm))

        # Transcript block — speaker-labelled if available, else plain
        tx_display = _spk_tx if _spk_tx else full_transcript
        if tx_display:
            story.append(Paragraph(f"ORIGINAL  [{lang_name} · {lang_code}]", tx_lbl_style))
            safe_tx = (tx_display
                       .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                       .replace("\n", "<br/>"))
            orig_tbl = Table([[Paragraph(safe_tx, mono_style)]], colWidths=[W])
            orig_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), colors.HexColor("#e8f0fb")),
                ("BOX",           (0,0), (-1,-1), 1, colors.HexColor("#90aad4")),
                ("LEFTPADDING",   (0,0), (-1,-1), 8),
                ("RIGHTPADDING",  (0,0), (-1,-1), 8),
                ("TOPPADDING",    (0,0), (-1,-1), 6),
                ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ]))
            story.append(orig_tbl)
            story.append(Spacer(1, 0.2*cm))

        # English translation
        if full_trans and full_trans != full_transcript and lang_code != "EN":
            story.append(Paragraph("ENGLISH TRANSLATION  [EN]", tx_en_style))
            safe_en = (full_trans
                       .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
            en_tbl = Table([[Paragraph(safe_en, mono_style)]], colWidths=[W])
            en_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,-1), colors.HexColor("#e8f5ee")),
                ("BOX",           (0,0), (-1,-1), 1, colors.HexColor("#90c8a4")),
                ("LEFTPADDING",   (0,0), (-1,-1), 8),
                ("RIGHTPADDING",  (0,0), (-1,-1), 8),
                ("TOPPADDING",    (0,0), (-1,-1), 6),
                ("BOTTOMPADDING", (0,0), (-1,-1), 6),
            ]))
            story.append(en_tbl)
        elif lang_code == "EN":
            story.append(Paragraph("ENGLISH TRANSLATION  [EN]", tx_en_style))
            story.append(Paragraph("Source language is English — no translation required.", mono_style))

    # ── Audio Preprocessing / SNR ──────────────────────────────────────────────
    _pre = result.get("preprocessing", {}) or {}
    if _pre.get("snr_before_db") is not None:
        story.append(_sec("AUDIO PREPROCESSING", section_style))
        story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
        story.append(Spacer(1, 0.15*cm))

        snr_b  = _pre["snr_before_db"]
        snr_a  = _pre.get("snr_after_db", snr_b)
        delta  = snr_a - snr_b
        sign   = "+" if delta >= 0 else ""
        snr_hdr_st = ParagraphStyle("SnrH", fontSize=8, textColor=WHITE,
                                     fontName="Helvetica-Bold", leading=11)
        snr_rows = [
            [Paragraph("METRIC",               snr_hdr_st),
             Paragraph("VALUE",                snr_hdr_st),
             Paragraph("NOTE",                 snr_hdr_st)],
            [Paragraph("SNR before denoising", label_style),
             Paragraph(f"{snr_b:.1f} dB",      body_style),
             Paragraph("Input audio quality",  body_style)],
            [Paragraph("SNR after denoising",  label_style),
             Paragraph(f"{snr_a:.1f} dB",      body_style),
             Paragraph(f"Improvement: {sign}{delta:.1f} dB", body_style)],
            [Paragraph("Duration",             label_style),
             Paragraph(f"{_pre.get('duration_sec', '?')}s", body_style),
             Paragraph("After preprocessing",  body_style)],
        ]
        snr_tbl = Table(snr_rows, colWidths=[4.5*cm, 3.5*cm, W-8*cm])
        snr_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
            ("TEXTCOLOR",     (0,0), (-1,0),  WHITE),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
            ("GRID",          (0,0), (-1,-1), 0.5, RULE),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 6),
        ]))
        story.append(snr_tbl)

    # ── Performance Metrics ────────────────────────────────────────────────────
    if metrics:
        story.append(_sec("PERFORMANCE METRICS", section_style))
        story.append(HRFlowable(width="100%", thickness=0.75, color=RULE))
        story.append(Spacer(1, 0.15*cm))

        m_lbl = ParagraphStyle("MLbl", fontSize=8, textColor=MUTED,
                                fontName="Helvetica-Bold", leading=11)
        m_val = ParagraphStyle("MVal", fontSize=9, textColor=DARK,
                                fontName="Helvetica", leading=13)
        m_hdr = ParagraphStyle("MHdr", fontSize=9, textColor=WHITE,
                                fontName="Helvetica-Bold", leading=12)

        # ── Tier 1: Auto metrics table ─────────────────────────────────────
        rtf = metrics.get("rtf", {})
        sc  = metrics.get("segment_confidence", {})
        ma  = metrics.get("model_agreement", {})
        ic  = metrics.get("isum_completeness", {})

        def _mv(val, fmt=None):
            if val is None:
                return Paragraph("N/A", m_val)
            return Paragraph(fmt.format(val) if fmt else str(val), m_val)

        auto_rows = [
            [Paragraph("METRIC", m_hdr), Paragraph("VALUE", m_hdr),
             Paragraph("DETAIL", m_hdr), Paragraph("GRADE", m_hdr)],
            [Paragraph("Real-Time Factor (RTF)", m_lbl),
             _mv(rtf.get("value"), "{:.3f}"),
             Paragraph(rtf.get("note", "-"), m_val),
             Paragraph(rtf.get("grade", "-"), m_val)],
            [Paragraph("Avg Segment Confidence", m_lbl),
             _mv(sc.get("mean"), "{:.3f}"),
             Paragraph(f"std={sc.get('std','-')}  low={sc.get('pct_low','-')}%  high={sc.get('pct_high','-')}%", m_val),
             Paragraph(sc.get("grade", "-"), m_val)],
            [Paragraph("No-Speech Probability", m_lbl),
             _mv(sc.get("mean_no_speech"), "{:.3f}"),
             Paragraph(f"{sc.get('count',0)} segments analysed", m_val),
             Paragraph("-", m_val)],
            [Paragraph("Model Agreement", m_lbl),
             Paragraph("AGREE" if ma.get("agree") else "DISAGREE" if ma.get("agree") is False else "N/A", m_val),
             Paragraph(f"Whisper={ma.get('whisper_lang','-').upper()} p={ma.get('whisper_prob',0):.3f}  "
                       f"FastText={ma.get('fasttext_lang','-').upper()} p={ma.get('fasttext_conf',0):.3f}", m_val),
             Paragraph(ma.get("grade", "-"), m_val)],
            [Paragraph("Ensemble LangID Score", m_lbl),
             _mv(ma.get("ensemble_score"), "{:.3f}"),
             Paragraph(f"Confidence delta: {ma.get('confidence_delta','-')}", m_val),
             Paragraph("-", m_val)],
            [Paragraph("5W ISUM Completeness", m_lbl),
             Paragraph(f"{ic.get('score',0)}/{ic.get('max',4)}  ({ic.get('pct',0)}%)", m_val),
             Paragraph("  ".join(f"{k.upper()}={'OK' if v else 'MISSING'}"
                                  for k, v in ic.get("fields", {}).items()), m_val),
             Paragraph(ic.get("grade", "-"), m_val)],
            [Paragraph("Keyword Density", m_lbl),
             Paragraph(f"{ic.get('kw_density',0)}%", m_val),
             Paragraph(f"{ic.get('kw_count',0)} alerts / {ic.get('word_count',0)} words", m_val),
             Paragraph("-", m_val)],
        ]

        col_w = [4.2*cm, 2.5*cm, W-9.5*cm, 2.8*cm]
        auto_tbl = Table(auto_rows, colWidths=col_w)
        auto_tbl.setStyle(TableStyle([
            ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
            ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
            ("GRID",          (0,0), (-1,-1), 0.5, RULE),
            ("VALIGN",        (0,0), (-1,-1), "TOP"),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 6),
            ("RIGHTPADDING",  (0,0), (-1,-1), 6),
        ]))
        story.append(auto_tbl)

        # ── Stage timings table ────────────────────────────────────────────
        st_data = metrics.get("stage_timings", {})
        if st_data.get("available") and st_data.get("timings"):
            story.append(Spacer(1, 0.3*cm))
            st_hdr = ParagraphStyle("STHdr", fontSize=8, textColor=GREEN,
                                     fontName="Helvetica-Bold", leading=11, spaceBefore=4)
            story.append(Paragraph("PIPELINE STAGE TIMINGS", st_hdr))
            story.append(Spacer(1, 0.1*cm))
            timing_rows = [[Paragraph("STAGE", m_hdr),
                            Paragraph("SECONDS", m_hdr),
                            Paragraph("% OF TOTAL", m_hdr)]]
            for stage, secs in st_data["timings"].items():
                pct = st_data.get("pcts", {}).get(stage, 0)
                timing_rows.append([
                    Paragraph(stage, m_lbl),
                    Paragraph(f"{secs}s", m_val),
                    Paragraph(f"{pct}%", m_val),
                ])
            t_tbl = Table(timing_rows, colWidths=[6*cm, 4*cm, W-10*cm])
            t_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
                ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
                ("GRID",          (0,0), (-1,-1), 0.5, RULE),
                ("TOPPADDING",    (0,0), (-1,-1), 4),
                ("BOTTOMPADDING", (0,0), (-1,-1), 4),
                ("LEFTPADDING",   (0,0), (-1,-1), 6),
            ]))
            story.append(t_tbl)

        # ── Memory / Vocab / Back-translation row ──────────────────────────
        mem  = metrics.get("memory", {})
        voc  = metrics.get("vocab_richness", {})
        bt   = metrics.get("backtrans_chrf")
        extra_rows = [[Paragraph("METRIC", m_hdr), Paragraph("VALUE", m_hdr),
                       Paragraph("DETAIL", m_hdr), Paragraph("GRADE", m_hdr)]]
        if mem.get("available"):
            extra_rows.append([
                Paragraph("Peak Memory Usage", m_lbl),
                Paragraph(f"{mem['peak_mb']} MB", m_val),
                Paragraph(f"start={mem['start_mb']} MB  delta=+{mem['delta_mb']} MB", m_val),
                Paragraph(mem.get("grade", "-"), m_val),
            ])
        if voc.get("ttr") is not None:
            extra_rows.append([
                Paragraph("Vocabulary Richness (TTR)", m_lbl),
                Paragraph(f"{voc['ttr']:.3f}", m_val),
                Paragraph(f"{voc['unique_words']} unique / {voc['word_count']} total words", m_val),
                Paragraph(voc.get("grade", "-"), m_val),
            ])
        if bt is not None:
            extra_rows.append([
                Paragraph("Back-Translation chrF", m_lbl),
                Paragraph(str(bt), m_val),
                Paragraph("EN -> source lang -> chrF vs original transcript", m_val),
                Paragraph("-", m_val),
            ])
        if len(extra_rows) > 1:
            story.append(Spacer(1, 0.15*cm))
            ex_tbl = Table(extra_rows, colWidths=col_w)
            ex_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
                ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
                ("GRID",          (0,0), (-1,-1), 0.5, RULE),
                ("VALIGN",        (0,0), (-1,-1), "TOP"),
                ("TOPPADDING",    (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
                ("LEFTPADDING",   (0,0), (-1,-1), 6),
                ("RIGHTPADDING",  (0,0), (-1,-1), 6),
            ]))
            story.append(ex_tbl)

        # ── Tier 2: Reference-based metrics (only if provided) ─────────────
        wer_r  = metrics.get("wer_result")
        bleu_r = metrics.get("bleu_result")

        if wer_r and not wer_r.get("error") and wer_r.get("wer") is not None:
            story.append(Spacer(1, 0.3*cm))
            ref_hdr_style = ParagraphStyle("RHdr", fontSize=8, textColor=GREEN,
                                            fontName="Helvetica-Bold", leading=11,
                                            spaceBefore=4)
            story.append(Paragraph("ASR QUALITY  —  WER / CER  (analyst-verified reference)", ref_hdr_style))
            story.append(Spacer(1, 0.1*cm))
            wer_rows = [
                [Paragraph("METRIC", m_hdr), Paragraph("VALUE", m_hdr),
                 Paragraph("BREAKDOWN", m_hdr), Paragraph("GRADE", m_hdr)],
                [Paragraph("Word Error Rate (WER)", m_lbl),
                 Paragraph(f"{wer_r['wer']}%", m_val),
                 Paragraph(f"S={wer_r.get('substitutions',0)}  "
                           f"D={wer_r.get('deletions',0)}  "
                           f"I={wer_r.get('insertions',0)}  "
                           f"ref={wer_r.get('ref_words',0)} words", m_val),
                 Paragraph(wer_r.get("grade", "-"), m_val)],
                [Paragraph("Character Error Rate (CER)", m_lbl),
                 Paragraph(f"{wer_r['cer']}%", m_val),
                 Paragraph(f"hyp={wer_r.get('hyp_words',0)} words", m_val),
                 Paragraph("-", m_val)],
            ]
            wer_tbl = Table(wer_rows, colWidths=col_w)
            wer_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
                ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
                ("GRID",          (0,0), (-1,-1), 0.5, RULE),
                ("VALIGN",        (0,0), (-1,-1), "TOP"),
                ("TOPPADDING",    (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
                ("LEFTPADDING",   (0,0), (-1,-1), 6),
                ("RIGHTPADDING",  (0,0), (-1,-1), 6),
            ]))
            story.append(wer_tbl)

        if bleu_r and not bleu_r.get("error") and bleu_r.get("bleu") is not None:
            story.append(Spacer(1, 0.3*cm))
            ref_hdr_style2 = ParagraphStyle("RHdr2", fontSize=8, textColor=GREEN,
                                             fontName="Helvetica-Bold", leading=11,
                                             spaceBefore=4)
            story.append(Paragraph("TRANSLATION QUALITY  —  BLEU / chrF / TER  (analyst-verified reference)", ref_hdr_style2))
            story.append(Spacer(1, 0.1*cm))
            prec = bleu_r.get("bleu_prec", [])
            bleu_rows = [
                [Paragraph("METRIC", m_hdr), Paragraph("VALUE", m_hdr),
                 Paragraph("DETAIL", m_hdr), Paragraph("GRADE", m_hdr)],
                [Paragraph("BLEU", m_lbl),
                 Paragraph(str(bleu_r["bleu"]), m_val),
                 Paragraph(f"BP={bleu_r.get('bleu_bp','-')}  "
                           f"n-gram: {' / '.join(str(p) for p in prec)}", m_val),
                 Paragraph(bleu_r.get("grade", "-"), m_val)],
                [Paragraph("chrF", m_lbl),
                 Paragraph(str(bleu_r["chrf"]), m_val),
                 Paragraph("Character n-gram F-score", m_val),
                 Paragraph("-", m_val)],
                [Paragraph("TER", m_lbl),
                 Paragraph(str(bleu_r["ter"]), m_val),
                 Paragraph("Translation Edit Rate (lower=better)", m_val),
                 Paragraph("-", m_val)],
            ]
            bleu_tbl = Table(bleu_rows, colWidths=col_w)
            bleu_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0), (-1,0),  HDR_BG),
                ("ROWBACKGROUNDS",(0,1), (-1,-1), [WHITE, ROW_ALT]),
                ("GRID",          (0,0), (-1,-1), 0.5, RULE),
                ("VALIGN",        (0,0), (-1,-1), "TOP"),
                ("TOPPADDING",    (0,0), (-1,-1), 5),
                ("BOTTOMPADDING", (0,0), (-1,-1), 5),
                ("LEFTPADDING",   (0,0), (-1,-1), 6),
                ("RIGHTPADDING",  (0,0), (-1,-1), 6),
            ]))
            story.append(bleu_tbl)

    # ── Footer ─────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.6*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=RULE))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(
        f"Generated by VANI Offline Intelligence System  |  UNCLASSIFIED  |  "
        f"{datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')}",
        footer_style,
    ))

    doc.build(story)
    return buf.getvalue()


def _lv(text: str, style) -> "Paragraph":
    from reportlab.platypus import Paragraph as P
    return P(str(text), style)


def _sec(text: str, style) -> "Paragraph":
    from reportlab.platypus import Paragraph as P
    return P(text.upper(), style)


# ── SRT subtitle export ────────────────────────────────────────────────────────

def build_srt(result: dict) -> bytes:
    """Generate SubRip (.srt) subtitle file from pipeline segments."""
    segments = result.get("segments", [])
    if not segments:
        dur  = result.get("isum", {}).get("processing_time_s", 5) or 5
        text = result.get("transcript", "").strip()
        return f"1\n00:00:00,000 --> 00:00:{int(dur):02d},000\n{text}\n".encode("utf-8")
    lines = []
    idx   = 1
    for seg in segments:
        text = seg.get("text", "").strip()
        if not text:
            continue
        spk  = seg.get("speaker", "")
        if spk:
            text = f"[{spk}] {text}"
        start = _srt_time(seg.get("start", 0.0))
        end   = _srt_time(seg.get("end",   seg.get("start", 0.0) + 1.0))
        lines.append(f"{idx}\n{start} --> {end}\n{text}")
        idx  += 1
    return "\n\n".join(lines).encode("utf-8")


def _srt_time(seconds: float) -> str:
    seconds = max(0.0, float(seconds))
    h  = int(seconds // 3600)
    m  = int((seconds % 3600) // 60)
    s  = int(seconds % 60)
    ms = int(round((seconds % 1) * 1000))
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


# ── CSV segment-level export (single intercept) ────────────────────────────────

def build_csv(result: dict) -> bytes:
    """CSV of timestamped segments for one intercept — useful for timeline analysis."""
    import csv
    segments  = result.get("segments", [])
    isum      = result.get("isum", {})
    report_id = isum.get("report_id", result.get("report_id", ""))
    lang      = result.get("final_language", "")
    buf       = io.StringIO()
    writer    = csv.writer(buf)
    writer.writerow(["report_id", "start_s", "end_s", "speaker",
                     "text", "confidence", "no_speech_prob", "language"])
    for seg in segments:
        conf    = seg.get("confidence")
        no_sp   = seg.get("no_speech_prob")
        writer.writerow([
            report_id,
            round(seg.get("start", 0.0), 3),
            round(seg.get("end",   0.0), 3),
            seg.get("speaker", ""),
            seg.get("text", "").strip(),
            round(conf,  4) if conf  is not None else "",
            round(no_sp, 4) if no_sp is not None else "",
            lang,
        ])
    return buf.getvalue().encode("utf-8")


# ── CSV bulk export (all intercepts from DB) ───────────────────────────────────

def build_bulk_csv(intercepts: list) -> bytes:
    """CSV summary of all intercepts as returned by db.get_all_intercepts()."""
    import csv
    buf    = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "report_id", "timestamp_utc", "audio_file", "language",
        "threat_level", "assessment", "who", "what", "where", "when",
        "categories", "flags",
    ])
    for row in intercepts:
        writer.writerow([
            row.get("report_id",       ""),
            row.get("timestamp_utc",   ""),
            row.get("audio_file",      ""),
            row.get("language",        ""),
            row.get("threat_level",    ""),
            row.get("assessment",      ""),
            row.get("who_field",       ""),
            row.get("what_field",      ""),
            row.get("where_field",     ""),
            row.get("when_field",      ""),
            row.get("categories",      ""),
            row.get("confidence_flags",""),
        ])
    return buf.getvalue().encode("utf-8")
